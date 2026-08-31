from __future__ import annotations

import base64
import os
import shutil
import subprocess
import tempfile
from datetime import date
from io import BytesIO
from pathlib import Path
from typing import Literal

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pydantic import BaseModel, ConfigDict, Field, model_validator

from cv_validator.ingestion.redaction import redact_national_ids_in_text


class _StrictModel(BaseModel):
    model_config = ConfigDict(extra="forbid")


class ProfileLink(_StrictModel):
    label: str
    url: str


class PersonalLinks(_StrictModel):
    linkedin: str | None = None
    github: str | None = None
    portfolio: str | None = None
    other: list[ProfileLink] = Field(default_factory=list)


class PersonalInformation(_StrictModel):
    first_name: str | None = None
    last_name: str | None = None
    email: str | None = None
    phone: str | None = None
    location: str | None = None
    links: PersonalLinks = Field(default_factory=PersonalLinks)


class ExperienceEntry(_StrictModel):
    id: str
    company: str | None = None
    company_category: str | None = None
    role: str | None = None
    project: str | None = None
    location: str | None = None
    start_date: str | None = None
    end_date: str | None = None
    current: bool = False
    responsibilities: list[str] = Field(default_factory=list)
    achievements: list[str] = Field(default_factory=list)
    technologies: list[str] = Field(default_factory=list)


class EducationEntry(_StrictModel):
    id: str
    institution: str | None = None
    degree: str | None = None
    field: str | None = None
    start_date: str | None = None
    end_date: str | None = None
    location: str | None = None
    description: str | None = None


class LanguageEntry(_StrictModel):
    id: str
    language: str
    level: str | None = None


class CertificationEntry(_StrictModel):
    id: str
    name: str
    issuer: str | None = None
    date: str | None = None
    url: str | None = None


class AdditionalSection(_StrictModel):
    id: str
    title: str
    items: list[str] = Field(default_factory=list)


CustomFieldKind = Literal["text", "number", "boolean", "date", "select"]
CustomFieldScalar = str | float | bool | None


class ProfileCustomFieldDefinition(_StrictModel):
    id: str = Field(min_length=1, max_length=80, pattern=r"^[A-Za-z0-9._-]+$")
    label: str = Field(min_length=1, max_length=120)
    kind: CustomFieldKind = "text"
    options: list[str] = Field(default_factory=list, max_length=30)
    default_value: CustomFieldScalar = None

    @model_validator(mode="after")
    def validate_options_and_default(self) -> ProfileCustomFieldDefinition:
        if redact_national_ids_in_text(self.id) != self.id:
            raise ValueError("custom field ID must not contain a national identifier")
        self.options = _normalize_custom_field_options(self.kind, self.options)
        _validate_custom_field_scalar(
            self.kind, self.default_value, self.options, field_name="default_value"
        )
        return self


class ProfileCustomFieldValue(_StrictModel):
    id: str = Field(min_length=1, max_length=80)
    label: str = Field(min_length=1, max_length=120)
    kind: CustomFieldKind = "text"
    value: CustomFieldScalar = None
    options: list[str] = Field(default_factory=list, max_length=30)

    @model_validator(mode="after")
    def validate_options_and_value(self) -> ProfileCustomFieldValue:
        if redact_national_ids_in_text(self.id) != self.id:
            raise ValueError("custom field ID must not contain a national identifier")
        self.options = _normalize_custom_field_options(self.kind, self.options)
        _validate_custom_field_scalar(
            self.kind, self.value, self.options, field_name="value"
        )
        return self


def _normalize_custom_field_options(
    kind: CustomFieldKind, options: list[str]
) -> list[str]:
    clean: list[str] = []
    seen: set[str] = set()
    for raw in options:
        option = raw.strip()
        if option and option not in seen:
            clean.append(option)
            seen.add(option)
    if kind == "select" and not clean:
        raise ValueError("select custom fields require at least one option")
    if kind != "select" and clean:
        raise ValueError("only select custom fields may define options")
    return clean


def _validate_custom_field_scalar(
    kind: CustomFieldKind,
    value: CustomFieldScalar,
    options: list[str],
    *,
    field_name: str,
) -> None:
    if value is None:
        return
    if kind == "boolean":
        if not isinstance(value, bool):
            raise ValueError(f"boolean custom field {field_name} must be boolean")
        return
    if kind == "number":
        if isinstance(value, bool) or not isinstance(value, (int, float)):
            raise ValueError(f"number custom field {field_name} must be numeric")
        return
    if not isinstance(value, str):
        raise ValueError(f"{kind} custom field {field_name} must be text")
    if kind == "select" and value not in options:
        raise ValueError(f"select custom field {field_name} must be one of its options")
    if kind == "date":
        try:
            date.fromisoformat(value)
        except ValueError as exc:
            raise ValueError(
                f"date custom field {field_name} must use YYYY-MM-DD"
            ) from exc


class CandidateProfile(_StrictModel):
    schema_version: Literal["candidate-profile-v1"] = "candidate-profile-v1"
    personal: PersonalInformation = Field(default_factory=PersonalInformation)
    headline: str | None = None
    summary: str | None = None
    skills: list[str] = Field(default_factory=list)
    technologies: list[str] = Field(default_factory=list)
    experience: list[ExperienceEntry] = Field(default_factory=list)
    education: list[EducationEntry] = Field(default_factory=list)
    languages: list[LanguageEntry] = Field(default_factory=list)
    certifications: list[CertificationEntry] = Field(default_factory=list)
    additional_sections: list[AdditionalSection] = Field(default_factory=list)
    custom_fields: list[ProfileCustomFieldValue] = Field(default_factory=list)


class AnonymizationPolicy(_StrictModel):
    hide_first_name: bool = False
    hide_last_name: bool = False
    hide_email: bool = False
    hide_phone: bool = False
    hide_location: bool = False
    hide_linkedin: bool = False
    hide_github: bool = False
    hide_portfolio: bool = False
    employer_mode: Literal["show", "hide", "genericize"] = "show"
    institution_mode: Literal["show", "hide"] = "show"


TemplateSectionKind = Literal[
    "summary",
    "skills",
    "technologies",
    "experience",
    "education",
    "languages",
    "certifications",
    "additional_sections",
    "custom_fields",
]
TemplateSectionLayout = Literal["default", "inline", "bullets"]


class ProfileTemplateBranding(_StrictModel):
    brand_name: str = Field(default="IDEGO", min_length=1, max_length=80)
    accent_hex: str = Field(default="#3CC2D9", pattern=r"^#[0-9A-Fa-f]{6}$")
    show_brand: bool = True


class ProfileTemplateTypography(_StrictModel):
    font_family: Literal["Aptos", "Arial", "Calibri"] = "Aptos"
    body_size: float = Field(default=10.5, ge=8, le=14)
    heading_size: float = Field(default=14, ge=10, le=22)


class ProfileTemplateHeader(_StrictModel):
    show_name: bool = True
    show_headline: bool = True
    show_contact: bool = True


class ProfileTemplateLogo(_StrictModel):
    data_url: str = Field(
        min_length=32,
        max_length=6_000_000,
        pattern=r"^data:image/png;base64,[A-Za-z0-9+/=]+$",
    )
    original_name: str = Field(min_length=1, max_length=255)
    x_pct: float = Field(default=72, ge=0, le=100)
    y_pct: float = Field(default=4, ge=0, le=100)
    width_pct: float = Field(default=18, ge=2, le=60)
    aspect_ratio: float = Field(default=2, gt=0.05, le=20)

    @model_validator(mode="after")
    def validate_page_bounds(self) -> ProfileTemplateLogo:
        if self.x_pct + self.width_pct > 100.001:
            raise ValueError("template logo must fit inside page width")
        height_pct = self.width_pct * (210 / 297) / self.aspect_ratio
        if self.y_pct + height_pct > 100.001:
            raise ValueError("template logo must fit inside page height")
        return self


class ProfileTemplateSection(_StrictModel):
    id: str = Field(min_length=1, max_length=80)
    kind: TemplateSectionKind
    title: str = Field(min_length=1, max_length=80)
    visible: bool = True
    layout: TemplateSectionLayout = "default"
    placement: Literal["full", "left", "right"] = "full"

    @model_validator(mode="after")
    def validate_safe_id(self) -> ProfileTemplateSection:
        if redact_national_ids_in_text(self.id) != self.id:
            raise ValueError("template section ID must not contain a national identifier")
        return self


class ProfileTemplate(_StrictModel):
    schema_version: Literal["profile-template-v1"] = "profile-template-v1"
    id: str = Field(min_length=1, max_length=96, pattern=r"^[A-Za-z0-9._-]+$")
    name: str = Field(min_length=1, max_length=120)
    visibility: Literal["private", "shared"] = "private"
    description: str | None = Field(default=None, max_length=300)
    branding: ProfileTemplateBranding = Field(default_factory=ProfileTemplateBranding)
    typography: ProfileTemplateTypography = Field(default_factory=ProfileTemplateTypography)
    header: ProfileTemplateHeader = Field(default_factory=ProfileTemplateHeader)
    logo: ProfileTemplateLogo | None = None
    sections: list[ProfileTemplateSection] = Field(min_length=1, max_length=9)

    @model_validator(mode="after")
    def validate_unique_sections(self) -> ProfileTemplate:
        if redact_national_ids_in_text(self.id) != self.id:
            raise ValueError("template ID must not contain a national identifier")
        ids = [section.id for section in self.sections]
        kinds = [section.kind for section in self.sections]
        if len(ids) != len(set(ids)):
            raise ValueError("template section IDs must be unique")
        if len(kinds) != len(set(kinds)):
            raise ValueError("template section kinds must be unique")
        return self


def _default_profile_builder_anonymization() -> AnonymizationPolicy:
    return AnonymizationPolicy(
        hide_first_name=True,
        hide_last_name=True,
        hide_email=True,
        hide_phone=True,
        hide_location=True,
        hide_linkedin=True,
        hide_github=True,
        hide_portfolio=True,
        employer_mode="hide",
        institution_mode="hide",
    )


class ProfileBuilderPreferences(_StrictModel):
    auto_summary: bool = False
    summary_instruction: str = Field(default="", max_length=12_000)
    anonymization: AnonymizationPolicy = Field(default_factory=_default_profile_builder_anonymization)
    aggregate_technologies: bool = True
    date_format: Literal["preserve", "yyyy-mm", "mm/yyyy", "yyyy"] = "preserve"
    default_template_id: str = Field(
        default="idego-default",
        min_length=1,
        max_length=96,
        pattern=r"^[A-Za-z0-9._-]+$",
    )
    filename_pattern: str = Field(default="{name}-profile", min_length=1, max_length=120)


class ProfileBuilderSnapshot(_StrictModel):
    source_filename: str = Field(min_length=1, max_length=512)
    profile: CandidateProfile
    anonymization: AnonymizationPolicy = Field(default_factory=AnonymizationPolicy)
    template: ProfileTemplate


class ProfileSummaryGenerationRequest(_StrictModel):
    profile: CandidateProfile
    instruction: str | None = Field(default=None, max_length=12_000)


ProfessionalSectionName = Literal[
    "headline",
    "summary",
    "skills",
    "technologies",
    "experience",
    "education",
    "languages",
    "certifications",
    "additional_sections",
]


class ProfessionalProfile(_StrictModel):
    headline: str | None = None
    summary: str | None = None
    skills: list[str] = Field(default_factory=list)
    technologies: list[str] = Field(default_factory=list)
    experience: list[ExperienceEntry] = Field(default_factory=list)
    education: list[EducationEntry] = Field(default_factory=list)
    languages: list[LanguageEntry] = Field(default_factory=list)
    certifications: list[CertificationEntry] = Field(default_factory=list)
    additional_sections: list[AdditionalSection] = Field(default_factory=list)


class ProfileTransformGenerationRequest(_StrictModel):
    profile: CandidateProfile
    sections: list[ProfessionalSectionName] = Field(min_length=1, max_length=10)
    instruction: str = Field(default="", max_length=12_000)
    mode: Literal["action", "translation"] = "action"
    target_language: Literal["en", "pl", "de", "fr", "es"] | None = None

    @model_validator(mode="after")
    def validate_transform(self) -> ProfileTransformGenerationRequest:
        if len(set(self.sections)) != len(self.sections):
            raise ValueError("transform sections must be unique")
        if self.mode == "translation" and self.target_language is None:
            raise ValueError("translation requires target language")
        return self


def professional_profile_from_candidate(profile: CandidateProfile) -> ProfessionalProfile:
    return ProfessionalProfile(
        headline=profile.headline,
        summary=profile.summary,
        skills=profile.skills,
        technologies=profile.technologies,
        experience=profile.experience,
        education=profile.education,
        languages=profile.languages,
        certifications=profile.certifications,
        additional_sections=profile.additional_sections,
    )


class ProfileExportRequest(_StrictModel):
    profile: CandidateProfile
    anonymization: AnonymizationPolicy = Field(default_factory=AnonymizationPolicy)
    template_id: str = "idego-default"
    template: ProfileTemplate | None = None

    @model_validator(mode="after")
    def validate_template_snapshot(self) -> ProfileExportRequest:
        if self.template is None and self.template_id != "idego-default":
            raise ValueError("custom template exports require an exact template snapshot")
        if self.template is not None and self.template.id != self.template_id:
            raise ValueError("template snapshot ID must match template_id")
        return self


IDEGO_NAVY = RGBColor(0x08, 0x19, 0x32)
MUTED = RGBColor(0x5C, 0x67, 0x73)


class ProfilePdfExportError(RuntimeError):
    """Safe PDF conversion failure without candidate content."""


def default_profile_template() -> ProfileTemplate:
    return ProfileTemplate(
        id="idego-default",
        name="IDEGO Default",
        visibility="shared",
        description="Default IDEGO candidate profile layout.",
        sections=[
            ProfileTemplateSection(id="summary", kind="summary", title="Summary"),
            ProfileTemplateSection(id="skills", kind="skills", title="Skills", layout="inline"),
            ProfileTemplateSection(
                id="technologies",
                kind="technologies",
                title="Technologies",
                layout="inline",
            ),
            ProfileTemplateSection(
                id="experience", kind="experience", title="Experience"
            ),
            ProfileTemplateSection(id="education", kind="education", title="Education"),
            ProfileTemplateSection(
                id="languages", kind="languages", title="Languages", layout="inline"
            ),
            ProfileTemplateSection(
                id="certifications",
                kind="certifications",
                title="Certifications",
                layout="bullets",
            ),
            ProfileTemplateSection(
                id="additional-sections",
                kind="additional_sections",
                title="Additional",
                layout="bullets",
            ),
            ProfileTemplateSection(
                id="custom-fields",
                kind="custom_fields",
                title="Details",
                layout="default",
            ),
        ],
    )


def sanitize_candidate_profile(profile: CandidateProfile) -> CandidateProfile:
    """Return a copy with supported national identifiers masked in every text field."""
    payload = _sanitize_profile_value(profile.model_dump(mode="python"))
    return CandidateProfile.model_validate(payload)


def sanitize_professional_profile(
    profile: ProfessionalProfile,
) -> ProfessionalProfile:
    payload = _sanitize_profile_value(profile.model_dump(mode="python"))
    return ProfessionalProfile.model_validate(payload)


def sanitize_profile_template(template: ProfileTemplate) -> ProfileTemplate:
    payload = template.model_dump(mode="python")
    payload["name"] = redact_national_ids_in_text(payload["name"])
    if payload["description"] is not None:
        payload["description"] = redact_national_ids_in_text(payload["description"])
    payload["branding"]["brand_name"] = redact_national_ids_in_text(
        payload["branding"]["brand_name"]
    )
    for section in payload["sections"]:
        section["title"] = redact_national_ids_in_text(section["title"])
    logo = payload.get("logo")
    if logo is not None:
        logo["original_name"] = sanitize_profile_builder_filename(
            logo["original_name"]
        )
    return ProfileTemplate.model_validate(payload)


def sanitize_profile_custom_field_definition(
    definition: ProfileCustomFieldDefinition,
) -> ProfileCustomFieldDefinition:
    payload = definition.model_dump(mode="python")
    payload["label"] = redact_national_ids_in_text(payload["label"])
    payload["options"] = [
        redact_national_ids_in_text(option) for option in payload["options"]
    ]
    if isinstance(payload["default_value"], str):
        payload["default_value"] = redact_national_ids_in_text(
            payload["default_value"]
        )
    return ProfileCustomFieldDefinition.model_validate(payload)


def sanitize_profile_builder_preferences(
    preferences: ProfileBuilderPreferences,
) -> ProfileBuilderPreferences:
    return preferences.model_copy(
        update={
            "summary_instruction": redact_national_ids_in_text(
                preferences.summary_instruction
            ),
            "filename_pattern": redact_national_ids_in_text(
                preferences.filename_pattern
            ),
        },
        deep=True,
    )


def sanitize_profile_builder_filename(filename: str) -> str:
    suffix = Path(filename).suffix
    stem = filename[:-len(suffix)] if suffix else filename
    return f"{redact_national_ids_in_text(stem)}{suffix}"


def sanitize_profile_builder_snapshot(
    snapshot: ProfileBuilderSnapshot,
) -> ProfileBuilderSnapshot:
    return snapshot.model_copy(
        update={
            "source_filename": sanitize_profile_builder_filename(
                snapshot.source_filename
            ),
            "profile": sanitize_candidate_profile(snapshot.profile),
            "template": sanitize_profile_template(snapshot.template),
        },
        deep=True,
    )


def _sanitize_profile_value(value):
    if isinstance(value, str):
        return redact_national_ids_in_text(value)
    if isinstance(value, list):
        return [_sanitize_profile_value(item) for item in value]
    if isinstance(value, dict):
        return {key: _sanitize_profile_value(item) for key, item in value.items()}
    return value


def materialize_custom_fields(
    profile: CandidateProfile,
    definitions: list[ProfileCustomFieldDefinition],
) -> CandidateProfile:
    result = profile.model_copy(deep=True)
    existing = {field.id: field for field in result.custom_fields}
    result.custom_fields = [
        existing.get(
            definition.id,
            ProfileCustomFieldValue(
                id=definition.id,
                label=definition.label,
                kind=definition.kind,
                value=definition.default_value,
                options=definition.options,
            ),
        )
        for definition in definitions
    ]
    return result


def apply_profile_conversion_preferences(
    profile: CandidateProfile,
    preferences: ProfileBuilderPreferences,
) -> CandidateProfile:
    result = profile.model_copy(deep=True)
    if preferences.aggregate_technologies:
        combined = [*result.technologies]
        for entry in result.experience:
            combined.extend(entry.technologies)
        result.technologies = _dedupe_strings(combined)
    if preferences.date_format != "preserve":
        for entry in result.experience:
            entry.start_date = _format_profile_date(entry.start_date, preferences.date_format)
            entry.end_date = _format_profile_date(entry.end_date, preferences.date_format)
        for entry in result.education:
            entry.start_date = _format_profile_date(entry.start_date, preferences.date_format)
            entry.end_date = _format_profile_date(entry.end_date, preferences.date_format)
        for entry in result.certifications:
            entry.date = _format_profile_date(entry.date, preferences.date_format)
    return result


def _format_profile_date(value: str | None, format_name: str) -> str | None:
    if not value:
        return value
    normalized = value.strip()
    year: str | None = None
    month: str | None = None
    import re
    match = re.fullmatch(r"(\d{4})[-/.](\d{1,2})(?:[-/.]\d{1,2})?", normalized)
    if match:
        year, month = match.group(1), match.group(2).zfill(2)
    else:
        match = re.fullmatch(r"(\d{1,2})[-/.](\d{4})", normalized)
        if match:
            month, year = match.group(1).zfill(2), match.group(2)
        elif re.fullmatch(r"\d{4}", normalized):
            year = normalized
    if year is None:
        return value
    if format_name == "yyyy":
        return year
    if month is None:
        return year
    if format_name == "yyyy-mm":
        return f"{year}-{month}"
    if format_name == "mm/yyyy":
        return f"{month}/{year}"
    return value


def _dedupe_strings(values: list[str]) -> list[str]:
    seen: set[str] = set()
    result: list[str] = []
    for raw in values:
        value = raw.strip()
        key = value.casefold()
        if value and key not in seen:
            seen.add(key)
            result.append(value)
    return result


def apply_profile_anonymization(
    profile: CandidateProfile,
    policy: AnonymizationPolicy,
) -> CandidateProfile:
    result = profile.model_copy(deep=True)
    if policy.hide_first_name:
        result.personal.first_name = None
    if policy.hide_last_name:
        result.personal.last_name = None
    if policy.hide_email:
        result.personal.email = None
    if policy.hide_phone:
        result.personal.phone = None
    if policy.hide_location:
        result.personal.location = None
    if policy.hide_linkedin:
        result.personal.links.linkedin = None
    if policy.hide_github:
        result.personal.links.github = None
    if policy.hide_portfolio:
        result.personal.links.portfolio = None
    if policy.employer_mode != "show":
        for item in result.experience:
            item.company = (
                None
                if policy.employer_mode == "hide"
                else item.company_category or "Company"
            )
    if policy.institution_mode == "hide":
        for item in result.education:
            item.institution = None
    return result


def render_candidate_profile_docx(
    profile: CandidateProfile,
    policy: AnonymizationPolicy,
    template: ProfileTemplate | None = None,
) -> bytes:
    selected_template = sanitize_profile_template(
        template or default_profile_template()
    )
    presentation = apply_profile_anonymization(
        sanitize_candidate_profile(profile), policy
    )
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    accent = _rgb_from_hex(selected_template.branding.accent_hex)
    font_family = selected_template.typography.font_family
    body_size = selected_template.typography.body_size
    heading_size = selected_template.typography.heading_size

    styles = document.styles
    styles["Normal"].font.name = font_family
    styles["Normal"].font.size = Pt(body_size)
    for style_name in ("Title", "Heading 1", "Heading 2"):
        styles[style_name].font.name = font_family
        styles[style_name].font.color.rgb = IDEGO_NAVY
    styles["Title"].font.size = Pt(min(heading_size + 10, 30))
    styles["Heading 1"].font.size = Pt(heading_size)
    styles["Heading 2"].font.size = Pt(min(body_size + 1.5, heading_size))

    header = section.header.paragraphs[0]
    if selected_template.branding.show_brand:
        header.text = selected_template.branding.brand_name
        header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        header.runs[0].font.bold = True
        header.runs[0].font.size = Pt(max(body_size + 2.5, 12))
        header.runs[0].font.color.rgb = accent
    if selected_template.logo is not None:
        _add_floating_template_logo(header, section, selected_template.logo)

    _render_profile_header(document, presentation, selected_template)
    _render_profile_template_sections(
        document,
        presentation,
        selected_template.sections,
        accent=accent,
    )

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_text = (
        f"Candidate Profile • {selected_template.branding.brand_name}"
        if selected_template.branding.show_brand
        else "Candidate Profile"
    )
    footer_run = footer.add_run(footer_text)
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = MUTED

    stream = BytesIO()
    document.save(stream)
    return stream.getvalue()


def render_candidate_profile_pdf(
    profile: CandidateProfile,
    policy: AnonymizationPolicy,
    template: ProfileTemplate | None = None,
    *,
    converter_path: str | None = None,
    timeout_seconds: float = 30.0,
) -> bytes:
    docx_bytes = render_candidate_profile_docx(profile, policy, template)
    converter = converter_path or shutil.which("libreoffice") or shutil.which("soffice")
    if not converter:
        raise ProfilePdfExportError("profile_pdf_converter_unavailable")

    try:
        with tempfile.TemporaryDirectory(prefix="profile-pdf-") as temp_dir:
            root = Path(temp_dir)
            source = root / "candidate-profile.docx"
            source.write_bytes(docx_bytes)
            profile_dir = root / "lo-profile"
            profile_dir.mkdir()
            command = [
                converter,
                "--headless",
                "--nologo",
                "--nodefault",
                "--nofirststartwizard",
                f"-env:UserInstallation={profile_dir.as_uri()}",
                "--convert-to",
                "pdf:writer_pdf_Export",
                "--outdir",
                str(root),
                str(source),
            ]
            env = os.environ.copy()
            env["HOME"] = str(root)
            completed = subprocess.run(
                command,
                stdin=subprocess.DEVNULL,
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                check=False,
                timeout=timeout_seconds,
                env=env,
            )
            output = root / "candidate-profile.pdf"
            if completed.returncode != 0 or not output.is_file():
                raise ProfilePdfExportError("profile_pdf_conversion_failed")
            pdf_bytes = output.read_bytes()
            if not pdf_bytes.startswith(b"%PDF-"):
                raise ProfilePdfExportError("profile_pdf_invalid_output")
            return pdf_bytes
    except subprocess.TimeoutExpired as exc:
        raise ProfilePdfExportError("profile_pdf_conversion_timeout") from exc
    except OSError as exc:
        raise ProfilePdfExportError("profile_pdf_conversion_failed") from exc


def _add_floating_template_logo(paragraph, section, logo: ProfileTemplateLogo) -> None:
    encoded = logo.data_url.split(",", 1)[1]
    try:
        image_bytes = base64.b64decode(encoded, validate=True)
    except (ValueError, base64.binascii.Error) as exc:
        raise ValueError("invalid template logo data") from exc

    page_width = int(section.page_width)
    page_height = int(section.page_height)
    width = max(1, int(page_width * logo.width_pct / 100))
    height = max(1, int(width / logo.aspect_ratio))
    x = int(page_width * logo.x_pct / 100)
    y = int(page_height * logo.y_pct / 100)

    run = paragraph.add_run()
    run.add_picture(BytesIO(image_bytes), width=width, height=height)
    drawing = run._r.drawing_lst[-1]
    inline = drawing[0]
    inline.tag = qn("wp:anchor")
    for name, value in {
        "distT": "0",
        "distB": "0",
        "distL": "0",
        "distR": "0",
        "simplePos": "0",
        "relativeHeight": "251658240",
        "behindDoc": "0",
        "locked": "0",
        "layoutInCell": "1",
        "allowOverlap": "1",
    }.items():
        inline.set(name, value)

    simple_pos = OxmlElement("wp:simplePos")
    simple_pos.set("x", "0")
    simple_pos.set("y", "0")
    position_h = OxmlElement("wp:positionH")
    position_h.set("relativeFrom", "page")
    pos_x = OxmlElement("wp:posOffset")
    pos_x.text = str(x)
    position_h.append(pos_x)
    position_v = OxmlElement("wp:positionV")
    position_v.set("relativeFrom", "page")
    pos_y = OxmlElement("wp:posOffset")
    pos_y.text = str(y)
    position_v.append(pos_y)

    inline.insert(0, simple_pos)
    inline.insert(1, position_h)
    inline.insert(2, position_v)

    extent_index = next(
        (index for index, child in enumerate(inline) if child.tag == qn("wp:extent")),
        3,
    )
    effect_extent = OxmlElement("wp:effectExtent")
    for name in ("l", "t", "r", "b"):
        effect_extent.set(name, "0")
    wrap_none = OxmlElement("wp:wrapNone")
    inline.insert(extent_index + 1, effect_extent)
    inline.insert(extent_index + 2, wrap_none)


def _render_profile_header(
    document: Document,
    presentation: CandidateProfile,
    template: ProfileTemplate,
) -> None:
    if template.header.show_name:
        name = " ".join(
            value
            for value in (
                presentation.personal.first_name,
                presentation.personal.last_name,
            )
            if value
        )
        title = document.add_paragraph(style="Title")
        title.paragraph_format.space_after = Pt(3)
        title_run = title.add_run(name or "Candidate Profile")
        title_run.font.color.rgb = IDEGO_NAVY

    if template.header.show_headline and presentation.headline:
        headline = document.add_paragraph()
        headline.paragraph_format.space_after = Pt(5)
        run = headline.add_run(presentation.headline)
        run.bold = True
        run.font.color.rgb = MUTED

    if template.header.show_contact:
        contact_values = [
            presentation.personal.email,
            presentation.personal.phone,
            presentation.personal.location,
            presentation.personal.links.linkedin,
            presentation.personal.links.github,
            presentation.personal.links.portfolio,
            *(
                f"{link.label}: {link.url}" if link.label else link.url
                for link in presentation.personal.links.other
            ),
        ]
        if any(contact_values):
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(8)
            run = paragraph.add_run(
                " • ".join(value for value in contact_values if value)
            )
            run.font.size = Pt(max(template.typography.body_size - 1.5, 8))
            run.font.color.rgb = MUTED


def _render_profile_template_sections(
    document: Document,
    presentation: CandidateProfile,
    sections: list[ProfileTemplateSection],
    *,
    accent: RGBColor,
) -> None:
    visible = [section for section in sections if section.visible]
    index = 0
    while index < len(visible):
        current = visible[index]
        if current.placement == "full":
            _render_profile_section(document, presentation, current, accent=accent)
            index += 1
            continue

        group: list[ProfileTemplateSection] = []
        while index < len(visible) and visible[index].placement != "full":
            group.append(visible[index])
            index += 1
        table = document.add_table(rows=1, cols=2)
        table.autofit = False
        row = table.rows[0]
        left_cell, right_cell = row.cells
        for section in group:
            target = left_cell if section.placement == "left" else right_cell
            _render_profile_section(target, presentation, section, accent=accent)
        for cell in (left_cell, right_cell):
            if cell.paragraphs and not cell.paragraphs[0].text and len(cell.paragraphs) > 1:
                first = cell.paragraphs[0]._element
                first.getparent().remove(first)
        document.add_paragraph().paragraph_format.space_after = Pt(0)


def _render_profile_section(
    document: Document,
    presentation: CandidateProfile,
    template_section: ProfileTemplateSection,
    *,
    accent: RGBColor,
) -> None:
    kind = template_section.kind
    if kind == "summary":
        if presentation.summary:
            _add_profile_heading(document, template_section.title, accent)
            document.add_paragraph(presentation.summary)
        return

    if kind == "skills":
        _render_string_values(
            document,
            template_section.title,
            presentation.skills,
            template_section.layout,
            accent,
        )
        return

    if kind == "technologies":
        _render_string_values(
            document,
            template_section.title,
            presentation.technologies,
            template_section.layout,
            accent,
        )
        return

    if kind == "experience":
        if not presentation.experience:
            return
        _add_profile_heading(document, template_section.title, accent)
        for entry in presentation.experience:
            heading_bits = [value for value in (entry.role, entry.company) if value]
            item_heading = document.add_paragraph(style="Heading 2")
            item_heading.paragraph_format.space_after = Pt(0)
            item_heading.add_run(" — ".join(heading_bits) or "Experience")
            _add_profile_meta_line(
                document,
                [
                    value
                    for value in (
                        _profile_date_range(
                            entry.start_date, entry.end_date, entry.current
                        ),
                        entry.location,
                        entry.project,
                    )
                    if value
                ],
            )
            for text in entry.responsibilities:
                document.add_paragraph(text, style="List Bullet")
            for text in entry.achievements:
                paragraph = document.add_paragraph(style="List Bullet")
                paragraph.add_run(text).bold = True
            if entry.technologies:
                paragraph = document.add_paragraph()
                paragraph.add_run("Technologies: ").bold = True
                paragraph.add_run(", ".join(entry.technologies))
        return

    if kind == "education":
        if not presentation.education:
            return
        _add_profile_heading(document, template_section.title, accent)
        for entry in presentation.education:
            title_bits = [value for value in (entry.degree, entry.field) if value]
            if entry.institution:
                title_bits.append(entry.institution)
            item_heading = document.add_paragraph(style="Heading 2")
            item_heading.paragraph_format.space_after = Pt(0)
            item_heading.add_run(" — ".join(title_bits) or "Education")
            _add_profile_meta_line(
                document,
                [
                    value
                    for value in (
                        _profile_date_range(entry.start_date, entry.end_date, False),
                        entry.location,
                    )
                    if value
                ],
            )
            if entry.description:
                document.add_paragraph(entry.description)
        return

    if kind == "languages":
        values = [
            " — ".join(value for value in (item.language, item.level) if value)
            for item in presentation.languages
        ]
        _render_string_values(
            document,
            template_section.title,
            values,
            template_section.layout,
            accent,
        )
        return

    if kind == "certifications":
        values: list[str] = []
        for certification in presentation.certifications:
            parts = [certification.name]
            if certification.issuer:
                parts.append(certification.issuer)
            if certification.date:
                parts.append(certification.date)
            if certification.url:
                parts.append(certification.url)
            values.append(" — ".join(parts))
        _render_string_values(
            document,
            template_section.title,
            values,
            template_section.layout,
            accent,
        )
        return

    if kind == "additional_sections":
        non_empty = [
            extra_section
            for extra_section in presentation.additional_sections
            if extra_section.items
        ]
        if not non_empty:
            return
        _add_profile_heading(document, template_section.title, accent)
        for extra_section in non_empty:
            item_heading = document.add_paragraph(style="Heading 2")
            item_heading.paragraph_format.space_after = Pt(0)
            item_heading.add_run(extra_section.title)
            for item in extra_section.items:
                document.add_paragraph(item, style="List Bullet")


    if kind == "custom_fields":
        values = [field for field in presentation.custom_fields if field.value not in (None, "")]
        if not values:
            return
        _add_profile_heading(document, template_section.title, accent)
        for field in values:
            paragraph = document.add_paragraph()
            paragraph.add_run(f"{field.label}: ").bold = True
            if isinstance(field.value, bool):
                paragraph.add_run("Yes" if field.value else "No")
            else:
                paragraph.add_run(str(field.value))


def _render_string_values(
    document: Document,
    title: str,
    values: list[str],
    layout: TemplateSectionLayout,
    accent: RGBColor,
) -> None:
    clean = [value for value in values if value]
    if not clean:
        return
    _add_profile_heading(document, title, accent)
    selected_layout = "inline" if layout == "default" else layout
    if selected_layout == "bullets":
        for value in clean:
            document.add_paragraph(value, style="List Bullet")
    else:
        document.add_paragraph(", ".join(clean))


def _add_profile_heading(
    document: Document,
    text: str,
    accent: RGBColor,
) -> None:
    paragraph = document.add_paragraph(style="Heading 1")
    paragraph.paragraph_format.space_before = Pt(8)
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(text)
    run.font.color.rgb = accent


def _add_profile_meta_line(document: Document, values: list[str]) -> None:
    if not values:
        return
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(2)
    run = paragraph.add_run(" • ".join(values))
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED


def _profile_date_range(
    start: str | None,
    end: str | None,
    current: bool,
) -> str | None:
    if current:
        end = "Present"
    if start and end:
        return f"{start} – {end}"
    return start or end


def _rgb_from_hex(value: str) -> RGBColor:
    normalized = value.lstrip("#")
    return RGBColor(
        int(normalized[0:2], 16),
        int(normalized[2:4], 16),
        int(normalized[4:6], 16),
    )
