from __future__ import annotations

import io
from collections.abc import Iterable

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

from cv_validator.config import IngestionConfig, load_ingestion_config
from cv_validator.file_links.extraction import (
    extract_docx_file_details,
    extract_docx_hyperlinks,
    merge_document_links,
)
from cv_validator.ingestion import IngestionError, PresentationSpan, RawDocument, SourceBlock, SourcePage
from cv_validator.structural.config import StructuralAuditConfig
from cv_validator.ingestion.text import validate_text_sufficiency


def extract_docx(
    content: bytes, config: IngestionConfig | None = None
) -> RawDocument:
    try:
        document = Document(io.BytesIO(content))
        page_texts = _extract_logical_pages(document.iter_inner_content())
    except Exception as exc:  # noqa: BLE001
        raise IngestionError(f"Failed to read DOCX: {exc}") from exc

    pages = tuple(
        SourcePage(
            page_id=f"page-{page_number:04d}",
            page_number=page_number,
            text=text,
        )
        for page_number, text in enumerate(page_texts, start=1)
    )
    try:
        presentation_spans, presentation_truncated = _presentation_spans(document, pages)
        source_blocks, source_blocks_partial = _source_blocks(document, pages)
        presentation_omitted = ()
    except Exception:  # Presentation inspection must not break text ingestion.
        presentation_spans, presentation_truncated = (), False
        source_blocks, source_blocks_partial = (), True
        presentation_omitted = ("docx_body_runs_unavailable",)
    omitted_parts = _omitted_parts(document) + presentation_omitted
    parsed = RawDocument(
        pages=pages,
        source_format="docx",
        file_details=extract_docx_file_details(document),
        document_links=merge_document_links(
            pages,
            extract_docx_hyperlinks(document, pages),
            source_format="docx",
        ),
        presentation_spans=presentation_spans,
        presentation_audited_parts=("docx_body_paragraph_runs", "docx_table_cell_runs", "docx_logical_page_breaks"),
        presentation_omitted_parts=omitted_parts,
        presentation_truncated=presentation_truncated,
        source_blocks=source_blocks,
        source_blocks_partial=source_blocks_partial,
    )
    validate_text_sufficiency(parsed, config or load_ingestion_config())
    return parsed


def _extract_logical_pages(blocks: Iterable[Paragraph | Table]) -> list[str]:
    page_fragments: list[list[str]] = [[]]

    for block in blocks:
        _append_block(block, page_fragments)

    return [
        _remove_paragraph_terminator("".join(fragments))
        for fragments in page_fragments
    ]


def _append_block(
    block: Paragraph | Table,
    page_fragments: list[list[str]],
) -> None:
    if isinstance(block, Paragraph):
        _append_paragraph(block, page_fragments)
        return

    seen_cells: set[int] = set()
    for row in block.rows:
        for cell in row.cells:
            cell_key = id(cell._tc)
            if cell_key in seen_cells:
                continue
            seen_cells.add(cell_key)
            for cell_block in cell.iter_inner_content():
                _append_block(cell_block, page_fragments)


def _append_paragraph(
    paragraph: Paragraph,
    page_fragments: list[list[str]],
) -> None:
    if _has_page_break_before(paragraph) and page_fragments[-1]:
        page_fragments.append([])

    for node in paragraph._p.iter():
        if node.tag == qn("w:t"):
            page_fragments[-1].append(node.text or "")
        elif node.tag == qn("w:tab"):
            page_fragments[-1].append("\t")
        elif node.tag == qn("w:br"):
            if node.get(qn("w:type")) == "page":
                page_fragments.append([])
            else:
                page_fragments[-1].append("\n")
        elif node.tag == qn("w:cr"):
            page_fragments[-1].append("\n")

    page_fragments[-1].append("\n")


def _has_page_break_before(paragraph: Paragraph) -> bool:
    properties = paragraph._p.pPr
    if properties is None:
        return False
    page_break = properties.find(qn("w:pageBreakBefore"))
    if page_break is None:
        return False
    return page_break.get(qn("w:val"), "true").lower() not in {
        "0",
        "false",
        "off",
    }


def _remove_paragraph_terminator(text: str) -> str:
    return text[:-1] if text.endswith("\n") else text


def _presentation_spans(document, pages):
    cfg = StructuralAuditConfig()
    spans = []
    page_index = 0
    offset = 0
    run_count = 0
    for block_index, block in enumerate(document.iter_inner_content()):
        paragraphs = [block] if isinstance(block, Paragraph) else [paragraph for row in block.rows for cell in row.cells for paragraph in cell.paragraphs]
        for paragraph_index, paragraph in enumerate(paragraphs):
            if _has_page_break_before(paragraph) and offset and page_index + 1 < len(pages):
                page_index += 1; offset = 0
            for run_index, run in enumerate(paragraph.runs):
                run_count += 1
                if run_count > cfg.max_docx_runs:
                    return tuple(spans), True
                text = run.text
                if not text:
                    continue
                page = pages[page_index]
                found = page.text.find(text, offset)
                exact = found >= 0
                if exact:
                    offset = found + len(text)
                color = run.font.color.rgb or (run.style.font.color.rgb if run.style is not None else None) or (paragraph.style.font.color.rgb if paragraph.style is not None else None)
                foreground = _rgb_luminance(str(color)) if color is not None else None
                shading = paragraph._p.pPr.find(qn("w:shd")) if paragraph._p.pPr is not None else None
                fill = shading.get(qn("w:fill")) if shading is not None else None
                background = _rgb_luminance(fill) if fill and fill.lower() != "auto" else None
                partial_offset = page.text.find(text[:1], offset) if not exact and text else -1
                association = "exact" if exact else "partial" if partial_offset >= 0 else "unmapped"
                spans.append(PresentationSpan(
                    page_id=page.page_id, page_number=page.page_number, text=text,
                    start_offset=found if exact else partial_offset if partial_offset >= 0 else None, end_offset=found + len(text) if exact else partial_offset + 1 if partial_offset >= 0 else None,
                    paragraph_path=f"body/{block_index}/paragraph/{paragraph_index}/run/{run_index}",
                    association=association,
                    font_size_points=_resolved_size(run, paragraph, document),
                    foreground_luminance=foreground, background_luminance=background,
                    explicit_hidden=bool(run.font.hidden or (run.style.font.hidden if run.style is not None else False) or (paragraph.style.font.hidden if paragraph.style is not None else False)),
                ))
    return tuple(spans), False


def _source_blocks(document, pages):
    blocks: list[SourceBlock] = []
    page_index = 0
    offset = 0

    def append_paragraph(paragraph, *, kind="paragraph", table_id=None, row_index=None, paragraph_path=None):
        nonlocal page_index, offset
        if _has_page_break_before(paragraph) and offset and page_index + 1 < len(pages):
            page_index += 1; offset = 0
        text = paragraph.text
        if not text:
            return
        page = pages[page_index]
        found = page.text.find(text, offset)
        exact = found >= 0
        if exact:
            offset = found + len(text)
        line_ids = tuple(line.line_id for line in page.lines if exact and line.start_offset < found + len(text) and line.end_offset >= found)
        num_pr = paragraph._p.pPr.numPr if paragraph._p.pPr is not None else None
        level = None
        if num_pr is not None and num_pr.ilvl is not None:
            try: level = int(num_pr.ilvl.val)
            except (TypeError, ValueError): level = None
        blocks.append(SourceBlock(
            id=f"source-block-{len(blocks)+1:04d}", page_id=page.page_id,
            page_number=page.page_number, source_order=len(blocks),
            kind="list_item" if num_pr is not None else kind, line_ids=line_ids,
            start_offset=found if exact else None, end_offset=found + len(text) if exact else None,
            paragraph_path=paragraph_path, table_id=table_id, row_index=row_index,
            list_level=level, association="exact" if exact else "unmapped",
        ))

    for block_index, block in enumerate(document.iter_inner_content()):
        if isinstance(block, Paragraph):
            append_paragraph(block, paragraph_path=f"body/{block_index}/paragraph")
            continue
        seen_cells: set[int] = set()
        for row_index, row in enumerate(block.rows):
            for cell_index, cell in enumerate(row.cells):
                if id(cell._tc) in seen_cells: continue
                seen_cells.add(id(cell._tc))
                for paragraph_index, paragraph in enumerate(cell.paragraphs):
                    append_paragraph(paragraph, kind="table_cell", table_id=f"table-{block_index:04d}", row_index=row_index, paragraph_path=f"body/{block_index}/row/{row_index}/cell/{cell_index}/paragraph/{paragraph_index}")
    return tuple(blocks), any(block.association != "exact" for block in blocks)


def _omitted_parts(document):
    omitted = set()
    if any(section.header.paragraphs and any(p.text for p in section.header.paragraphs) for section in document.sections): omitted.add("docx_headers")
    if any(section.footer.paragraphs and any(p.text for p in section.footer.paragraphs) for section in document.sections): omitted.add("docx_footers")
    xml = document.part.element.xml
    for marker, name in (("w:txbxContent", "docx_textboxes"), ("w:commentReference", "docx_comments"), ("w:drawing", "docx_drawings"), ("w:footnoteReference", "docx_footnotes"), ("w:endnoteReference", "docx_endnotes"), ("w:object", "docx_embedded_files")):
        if marker in xml: omitted.add(name)
    return tuple(sorted(omitted))


def _rgb_luminance(value):
    if not value or len(value) != 6:
        return None
    try: rgb = [int(value[index:index+2], 16) / 255 for index in (0, 2, 4)]
    except ValueError: return None
    return 0.2126 * rgb[0] + 0.7152 * rgb[1] + 0.0722 * rgb[2]


def _resolved_size(run, paragraph, document):
    for size in (run.font.size, run.style.font.size if run.style is not None else None, paragraph.style.font.size if paragraph.style is not None else None, document.styles["Normal"].font.size):
        if size is not None:
            return float(size.pt)
    return None
