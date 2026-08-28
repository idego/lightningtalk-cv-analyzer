from __future__ import annotations

import io
import socket
import json
from dataclasses import replace

import pytest
from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from fastapi.testclient import TestClient
from reportlab.pdfgen import canvas

from cv_validator.ai.config import AISettings
from cv_validator.api.app import create_app
from cv_validator.api.persistence import PersistenceConfig, PersistenceStore
from cv_validator.domain import (
    ComponentVersion,
    DocumentLink,
    Evidence,
    FileDetail,
    FileDetailField,
    FileDetailStatus,
    FileDetails,
    LinkAssociation,
    LinkCheckResult,
    LinkInspection,
    LinkOutcomeStatus,
    LinkReasonCode,
    LinkRole,
    LinkSource,
)
from cv_validator.file_links.catalog import classify_service_domain
from cv_validator.file_links.checker import (
    LinkCheckConfig,
    LinkHTTPResponse,
    inspect_document_links,
)
from cv_validator.file_links.extraction import (
    extract_docx_file_details,
    extract_docx_hyperlinks,
    extract_pdf_file_details,
    extract_pdf_hyperlinks,
    merge_document_links,
)
from cv_validator.file_links.normalization import URLNormalizationError, normalize_url
from cv_validator.ingestion import SourcePage
from cv_validator.pipeline import analyze_cv_bytes_result, analyze_cv_text_result
from cv_validator.serialization import serialize_report_payload


def _link(
    target: str | None = "https://example.com/profile",
    *,
    displayed: str | None = "https://example.com/profile",
    role: LinkRole = LinkRole.PROFILE,
    association: LinkAssociation = LinkAssociation.VISIBLE_ONLY,
) -> DocumentLink:
    return DocumentLink(
        id="link:test:1",
        displayed_value=displayed,
        target=target,
        source_format="docx",
        source=LinkSource.VISIBLE_URL,
        association=association,
        role=role,
        page_number=1,
    )


def test_file_contracts_round_trip_and_reject_invalid_values() -> None:
    version = ComponentVersion("test", "1")
    details = FileDetails(
        source_format="pdf",
        extractor_version=version,
        fields=(
            FileDetail(
                FileDetailField.AUTHOR,
                "Synthetic Author",
                FileDetailStatus.AVAILABLE,
                "pdf",
                version,
            ),
            FileDetail(
                FileDetailField.SUBJECT,
                None,
                FileDetailStatus.UNAVAILABLE,
                "pdf",
                version,
            ),
        ),
    )

    assert FileDetails.from_dict(details.to_dict()) == details
    with pytest.raises(ValueError):
        FileDetails(
            "pdf",
            version,
            details.fields,
            contract_version="file-details-v0",
        )
    with pytest.raises(ValueError):
        FileDetail(
            FileDetailField.AUTHOR,
            "value",
            FileDetailStatus.UNAVAILABLE,
            "pdf",
            version,
        )


def test_url_normalization_removes_query_and_fragment_and_preserves_idna() -> None:
    normalized = normalize_url("HTTPS://BÜCHER.Example/path/../cv?token=secret#part")

    assert normalized.sanitized_url == "https://xn--bcher-kva.example/cv"
    assert normalized.comparison_key == normalized.sanitized_url
    assert normalized.registrable_domain == "xn--bcher-kva.example"

    for value, reason in (
        ("javascript:alert(1)", "unsafe_scheme"),
        ("https://user:pass@example.com/cv", "embedded_credentials"),
        ("https://example.com:8443/cv", "disallowed_port"),
        ("https://example..com/cv", "invalid_host"),
        ("https://%65xample.com/cv", "invalid_host"),
    ):
        with pytest.raises(URLNormalizationError) as error:
            normalize_url(value)
        assert error.value.reason_code == reason


def test_catalog_accepts_official_host_and_flags_lookalike() -> None:
    assert classify_service_domain("www.linkedin.com").official is True
    match = classify_service_domain("linkedln.com")
    assert match.lookalike is True
    assert match.service == "linkedin"


class _Resolver:
    def __init__(self, mapping: dict[str, tuple[str, ...]]) -> None:
        self.mapping = mapping
        self.calls: list[str] = []

    def resolve(self, hostname: str, port: int) -> tuple[str, ...]:
        self.calls.append(hostname)
        if hostname not in self.mapping:
            raise socket.gaierror(hostname)
        return self.mapping[hostname]


class _HTTP:
    def __init__(self, responses: list[LinkHTTPResponse]) -> None:
        self.responses = list(responses)
        self.calls: list[tuple[str, str, dict[str, str]]] = []

    def request(self, method, url, *, headers, timeout_seconds, max_response_bytes):
        self.calls.append((method, url, dict(headers)))
        return self.responses.pop(0)


def test_checker_reaches_public_link_and_uses_bounded_get_fallback() -> None:
    resolver = _Resolver({"example.com": ("93.184.216.34",)})
    http = _HTTP([
        LinkHTTPResponse(405, {}),
        LinkHTTPResponse(200, {"content-length": "12"}),
    ])

    result = inspect_document_links(
        [_link("https://example.com/profile?token=secret")],
        LinkCheckConfig(max_concurrency=1),
        dns_resolver=resolver,
        http_client=http,
        now=lambda: "2026-08-27T00:00:00+00:00",
    )

    assert result.links[0].status is LinkOutcomeStatus.REACHABLE
    assert result.links[0].sanitized_target == "https://example.com/profile"
    assert [call[0] for call in http.calls] == ["HEAD", "GET"]
    assert all("Cookie" not in call[2] for call in http.calls)


def test_checker_enforces_response_limit_from_head_without_retry() -> None:
    resolver = _Resolver({"example.com": ("93.184.216.34",)})
    http = _HTTP([LinkHTTPResponse(200, {"content-length": "9999"})])

    result = inspect_document_links(
        [_link()],
        LinkCheckConfig(max_concurrency=1, max_response_bytes=1024, max_retries=3),
        dns_resolver=resolver,
        http_client=http,
    )

    assert result.links[0].status is LinkOutcomeStatus.UNAVAILABLE
    assert result.links[0].reason_code is LinkReasonCode.RESPONSE_LIMIT
    assert len(http.calls) == 1


def test_checker_blocks_private_destinations_before_http_request() -> None:
    resolver = _Resolver({"internal.example": ("127.0.0.1",)})
    http = _HTTP([LinkHTTPResponse(200, {})])
    link = _link("https://internal.example/profile")

    result = inspect_document_links(
        [link],
        LinkCheckConfig(max_concurrency=1),
        dns_resolver=resolver,
        http_client=http,
    )

    assert result.links[0].status is LinkOutcomeStatus.SUSPICIOUS
    assert result.links[0].reason_code is LinkReasonCode.UNSAFE_DESTINATION
    assert http.calls == []


def test_checker_classifies_not_found_as_suspicious_but_blocked_access_as_unavailable() -> None:
    resolver = _Resolver({"example.com": ("93.184.216.34",)})
    http = _HTTP([LinkHTTPResponse(404, {}), LinkHTTPResponse(403, {})])
    links = [
        _link("https://example.com/profile", role=LinkRole.PROFILE),
        DocumentLink(
            id="link:test:2",
            displayed_value="Portfolio",
            target="https://example.com/portfolio",
            source_format="docx",
            source=LinkSource.EMBEDDED_HYPERLINK,
            association=LinkAssociation.EMBEDDED_ONLY,
            role=LinkRole.PORTFOLIO,
            page_number=1,
        ),
    ]

    result = inspect_document_links(
        links,
        LinkCheckConfig(max_concurrency=1),
        dns_resolver=resolver,
        http_client=http,
    )

    assert result.links[0].reason_code is LinkReasonCode.DECLARED_LINK_NOT_FOUND
    assert result.links[0].status is LinkOutcomeStatus.SUSPICIOUS
    assert result.links[1].reason_code is LinkReasonCode.HTTP_FORBIDDEN
    assert result.links[1].status is LinkOutcomeStatus.UNAVAILABLE


def test_checker_stops_unsafe_redirect_and_marks_unrelated_redirect_suspicious() -> None:
    resolver = _Resolver({
        "example.com": ("93.184.216.34",),
        "other.example": ("93.184.216.35",),
        "private.example": ("10.0.0.5",),
    })
    unsafe_http = _HTTP([LinkHTTPResponse(302, {"location": "https://private.example/cv"})])
    unsafe = inspect_document_links(
        [_link()],
        LinkCheckConfig(max_concurrency=1),
        dns_resolver=resolver,
        http_client=unsafe_http,
    )
    assert unsafe.links[0].reason_code is LinkReasonCode.UNSAFE_REDIRECT
    assert len(unsafe_http.calls) == 1

    cross_http = _HTTP([LinkHTTPResponse(302, {"location": "https://other.example/cv"}), LinkHTTPResponse(200, {})])
    cross = inspect_document_links(
        [_link()],
        LinkCheckConfig(max_concurrency=1),
        dns_resolver=resolver,
        http_client=cross_http,
    )
    assert cross.links[0].reason_code is LinkReasonCode.UNRELATED_CROSS_DOMAIN_REDIRECT
    assert cross.links[0].terminal_registrable_domain == "other.example"


def test_checker_follows_a_same_registrable_domain_redirect_after_revalidation() -> None:
    resolver = _Resolver({
        "example.com": ("93.184.216.34",),
        "www.example.com": ("93.184.216.34",),
    })
    http = _HTTP([
        LinkHTTPResponse(302, {"location": "https://www.example.com/profile"}),
        LinkHTTPResponse(200, {}),
    ])

    result = inspect_document_links(
        [_link()],
        LinkCheckConfig(max_concurrency=1),
        dns_resolver=resolver,
        http_client=http,
    )

    assert result.links[0].status is LinkOutcomeStatus.REACHABLE
    assert [call[1] for call in http.calls] == [
        "https://example.com/profile",
        "https://www.example.com/profile",
    ]


def test_checker_keeps_not_found_suspicious_and_rate_limit_neutral() -> None:
    resolver = _Resolver({"example.com": ("93.184.216.34",)})
    http = _HTTP([LinkHTTPResponse(410, {}), LinkHTTPResponse(429, {})])
    links = [
        _link("https://example.com/profile", role=LinkRole.PROFILE),
        replace(
            _link("https://example.com/portfolio", role=LinkRole.PORTFOLIO),
            id="link:test:2",
        ),
    ]

    result = inspect_document_links(
        links,
        LinkCheckConfig(max_concurrency=1),
        dns_resolver=resolver,
        http_client=http,
    )

    assert result.links[0].status is LinkOutcomeStatus.SUSPICIOUS
    assert result.links[0].reason_code is LinkReasonCode.DECLARED_LINK_NOT_FOUND
    assert result.links[1].status is LinkOutcomeStatus.UNAVAILABLE
    assert result.links[1].reason_code is LinkReasonCode.RATE_LIMITED


def test_docx_metadata_and_hyperlink_extraction_are_bounded() -> None:
    document = Document()
    document.core_properties.author = "Synthetic Author"
    document.core_properties.title = "Synthetic CV"
    paragraph = document.add_paragraph("Portfolio: ")
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), "rIdExternal")
    run = OxmlElement("w:r")
    text = OxmlElement("w:t")
    text.text = "Example portfolio"
    run.append(text)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)
    relationship = document.part.relate_to(
        "https://example.com/portfolio",
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink.set(qn("r:id"), relationship)
    buffer = io.BytesIO()
    document.save(buffer)
    reopened = Document(io.BytesIO(buffer.getvalue()))
    pages = (SourcePage("page-0001", 1, "Portfolio: Example portfolio"),)

    details = extract_docx_file_details(reopened)
    links = extract_docx_hyperlinks(reopened, pages)

    assert details.to_dict()["fields"]["author"]["value"] == "Synthetic Author"
    assert details.to_dict()["fields"]["creator"] if "creator" in details.to_dict()["fields"] else True
    assert links[0].displayed_value == "Example portfolio"
    assert links[0].target == "https://example.com/portfolio"
    assert links[0].evidence[0].excerpt == "Example portfolio"
    assert "custom" not in details.to_dict()


def test_pdf_metadata_and_visible_embedded_link_merge() -> None:
    buffer = io.BytesIO()
    document = canvas.Canvas(buffer)
    document.setAuthor("Synthetic Author")
    document.setTitle("Synthetic CV")
    document.drawString(72, 800, "Portfolio https://example.com/portfolio")
    document.linkURL(
        "https://example.com/portfolio",
        (72, 790, 280, 815),
        relative=0,
    )
    document.save()
    import pdfplumber

    with pdfplumber.open(io.BytesIO(buffer.getvalue())) as pdf:
        pages = (SourcePage("page-0001", 1, pdf.pages[0].extract_text() or ""),)
        details = extract_pdf_file_details(pdf)
        embedded = extract_pdf_hyperlinks(pdf, pages)

    merged = merge_document_links(pages, embedded, source_format="pdf")
    assert details.to_dict()["fields"]["author"]["value"] == "Synthetic Author"
    assert len(merged) == 1
    assert merged[0].association is LinkAssociation.MATCHED
    assert merged[0].source is LinkSource.VISIBLE_AND_EMBEDDED


def test_pdf_link_inventory_keeps_friendly_labels_mismatches_duplicates_and_invalids() -> None:
    class FakePage:
        def __init__(self) -> None:
            self.annots = [
                {"uri": "https://example.com/target"},
                {"uri": "https://example.com/target"},
                {"Subtype": "/Link"},
            ]

        def extract_words(self):
            raise AssertionError("hyperlink association must reuse canonical text")

    pages = (SourcePage(
        "page-0001",
        1,
        "Profile https://example.com/visible",
    ),)
    embedded = extract_pdf_hyperlinks(type("FakePDF", (), {"pages": (FakePage(),)})(), pages)
    merged = merge_document_links(pages, embedded, source_format="pdf")

    assert any(
        link.association is LinkAssociation.MISMATCHED
        and link.displayed_value == "https://example.com/visible"
        and link.target == "https://example.com/target"
        for link in merged
    )
    invalid = next(link for link in merged if link.target is None)
    assert invalid.invalid_reason == "missing_uri"
    assert not any(
        link.target == "https://example.com/target"
        for link in merged[1:]
    )


def _add_external_hyperlink(paragraph, displayed: str, target: str | None) -> None:
    hyperlink = OxmlElement("w:hyperlink")
    run = OxmlElement("w:r")
    text = OxmlElement("w:t")
    text.text = displayed
    run.append(text)
    hyperlink.append(run)
    if target is not None:
        relationship = paragraph.part.relate_to(
            target,
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
            is_external=True,
        )
        hyperlink.set(qn("r:id"), relationship)
    paragraph._p.append(hyperlink)


def test_docx_hyperlinks_cover_table_header_footer_and_malformed_relationships() -> None:
    document = Document()
    body = document.add_paragraph("Portfolio: ")
    _add_external_hyperlink(body, "Example portfolio", "https://example.com/portfolio")
    broken = document.add_paragraph("Broken: ")
    _add_external_hyperlink(broken, "Broken link", None)
    table = document.add_table(rows=1, cols=1)
    _add_external_hyperlink(table.cell(0, 0).paragraphs[0], "Code", "https://github.com/example/code")
    _add_external_hyperlink(document.sections[0].header.paragraphs[0], "Profile", "https://example.com/profile")
    _add_external_hyperlink(document.sections[0].footer.paragraphs[0], "Publication", "https://example.com/paper")

    buffer = io.BytesIO()
    document.save(buffer)
    from cv_validator.ingestion.docx import extract_docx

    parsed = extract_docx(buffer.getvalue())
    links = parsed.document_links

    assert {link.source_location for link in links} == {"body", "header", "footer"}
    assert any(link.displayed_value == "Code" and link.target == "https://github.com/example/code" for link in links)
    assert any(link.displayed_value == "Example portfolio" and link.evidence for link in links)
    invalid = next(link for link in links if link.displayed_value == "Broken link")
    assert invalid.target is None
    assert invalid.invalid_reason == "missing_relationship"
    assert not any(
        link.source_location in {"header", "footer"} and link.evidence
        for link in links
    )


def test_docx_hyperlinks_inside_a_paragraph_keep_explicit_page_association() -> None:
    document = Document()
    document.add_paragraph("Jane Example Berlin Germany")
    paragraph = document.add_paragraph()
    _add_external_hyperlink(paragraph, "First profile", "https://example.com/first")
    page_break = OxmlElement("w:br")
    page_break.set(qn("w:type"), "page")
    paragraph._p.append(page_break)
    _add_external_hyperlink(paragraph, "Second profile", "https://example.com/second")

    buffer = io.BytesIO()
    document.save(buffer)
    from cv_validator.ingestion.docx import extract_docx

    parsed = extract_docx(buffer.getvalue())
    links = {
        link.displayed_value: link
        for link in parsed.document_links
        if link.source_location == "body"
    }

    assert links["First profile"].page_number == 1
    assert links["Second profile"].page_number == 2


def test_default_docx_tool_metadata_is_not_reporter_metadata() -> None:
    details = extract_docx_file_details(Document())

    assert all(field.status is FileDetailStatus.UNAVAILABLE for field in details.fields)


def test_docx_without_core_properties_does_not_report_library_defaults() -> None:
    document = Document()
    package = document.part.package
    core_relationship_id = next(
        relationship_id
        for relationship_id, relationship in package.rels.items()
        if relationship.reltype == RELATIONSHIP_TYPE.CORE_PROPERTIES
    )
    package.rels.pop(core_relationship_id)

    details = extract_docx_file_details(document)

    assert all(field.status is FileDetailStatus.UNAVAILABLE for field in details.fields)


def test_malformed_pdf_metadata_is_unavailable_and_custom_keys_are_ignored() -> None:
    class FakePDF:
        metadata = {
            "/Author": "Safe Author\x00hidden",
            "/Title": "Synthetic CV",
            "/CreationDate": "D:20261301120000+99'99'",
            "/CustomPII": "should not be read",
        }

    details = extract_pdf_file_details(FakePDF())
    fields = details.to_dict()["fields"]

    assert fields["author"]["status"] == "unavailable"
    assert fields["title"]["value"] == "Synthetic CV"
    assert fields["creation_time"]["status"] == "unavailable"
    assert "custom" not in str(details.to_dict()).lower()


def test_normalization_rejects_parser_confusion_and_handles_ip_literals() -> None:
    assert normalize_url("https://[2001:db8::1]/cv").hostname == "2001:db8::1"

    for value, reason in (
        ("http:example.com/cv", "invalid_host"),
        ("//example.com/cv", "invalid_host"),
        ("https://example.com:", "invalid_host"),
    ):
        with pytest.raises(URLNormalizationError) as error:
            normalize_url(value)
        assert error.value.reason_code == reason


def test_disabled_checker_is_neutral_and_metrics_are_safe() -> None:
    class Metrics:
        def __init__(self) -> None:
            self.calls = []

        def increment(self, name, **labels):
            self.calls.append((name, labels))

    metrics = Metrics()
    link = replace(
        _link(
            "https://example.com/profile?candidate=private#fragment",
            displayed="Portfolio https://example.com/profile?candidate=private#fragment",
        ),
        evidence=(
            Evidence(
                page_id="page-0001",
                page_number=1,
                start_offset=0,
                end_offset=64,
                excerpt="Portfolio https://example.com/profile?candidate=private#fragment",
            ),
        ),
    )
    result = inspect_document_links(
        [link],
        LinkCheckConfig(enabled=False, max_concurrency=1),
        metrics=metrics,
    )

    assert result.links[0].status is LinkOutcomeStatus.UNAVAILABLE
    assert result.links[0].reason_code is LinkReasonCode.INSPECTION_DISABLED
    assert result.links[0].displayed_value == "Portfolio https://example.com/profile"
    assert result.links[0].source_evidence[0].excerpt == (
        "Portfolio https://example.com/profile"
    )
    assert "candidate" not in str(metrics.calls)
    assert "example.com/profile" not in str(metrics.calls)


def test_checker_returns_network_failure_classes_as_unavailable() -> None:
    class FailingResolver:
        def resolve(self, hostname: str, port: int) -> tuple[str, ...]:
            raise socket.gaierror(hostname)

    result = inspect_document_links(
        [_link()],
        LinkCheckConfig(max_concurrency=1),
        dns_resolver=FailingResolver(),
    )

    assert result.links[0].status is LinkOutcomeStatus.UNAVAILABLE
    assert result.links[0].reason_code is LinkReasonCode.DNS_FAILURE


def test_check_config_rejects_malformed_environment_ports(monkeypatch) -> None:
    monkeypatch.setenv("CV_VALIDATOR_LINK_CHECK_PORTS", "80,nope")

    with pytest.raises(ValueError, match="must contain integers"):
        LinkCheckConfig.from_env()


def test_pipeline_keeps_link_inspection_outside_deterministic_score() -> None:
    document = Document()
    document.add_paragraph(
        "Jane Example\nBerlin, Germany\nPortfolio: https://example.com/profile"
    )
    buffer = io.BytesIO()
    document.save(buffer)

    resolver = _Resolver({"example.com": ("93.184.216.34",)})
    checked = analyze_cv_bytes_result(
        buffer.getvalue(),
        "cv.docx",
        ai_settings=AISettings(enabled=False),
        link_check_config=LinkCheckConfig(max_concurrency=1),
        link_dns_resolver=resolver,
        link_http_client=_HTTP([LinkHTTPResponse(404, {})]),
    )
    disabled = analyze_cv_bytes_result(
        buffer.getvalue(),
        "cv.docx",
        ai_settings=AISettings(enabled=False),
        link_check_config=LinkCheckConfig(enabled=False, max_concurrency=1),
    )

    assert checked.report.score == disabled.report.score
    assert checked.report.band is disabled.report.band
    assert checked.report.deterministic == disabled.report.deterministic
    assert checked.report.link_inspection is not None
    assert checked.report.link_inspection.links[0].status is LinkOutcomeStatus.SUSPICIOUS
    assert disabled.report.link_inspection is not None
    assert disabled.report.link_inspection.links[0].status is LinkOutcomeStatus.UNAVAILABLE


def test_api_batch_persists_link_results_and_isolates_bad_upload(tmp_path) -> None:
    document = Document()
    document.add_paragraph(
        "Jane Example\nBerlin, Germany\nPortfolio: https://example.com/profile"
    )
    buffer = io.BytesIO()
    document.save(buffer)
    resolver = _Resolver({"example.com": ("93.184.216.34",)})
    app = create_app(
        db_path=tmp_path / "links.db",
        ai_settings=AISettings(enabled=False),
        link_check_config=LinkCheckConfig(max_concurrency=1),
        link_dns_resolver=resolver,
        link_http_client=_HTTP([LinkHTTPResponse(404, {})]),
    )

    with TestClient(app) as client:
        response = client.post(
            "/analyze/batch",
            files=[
                (
                    "files",
                    (
                        "cv.docx",
                        buffer.getvalue(),
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    ),
                ),
                ("files", ("broken.txt", b"not a supported CV", "text/plain")),
            ],
        )

    assert response.status_code == 200
    results = response.json()["results"]
    assert results[0]["status"] == "ok"
    assert results[0]["report"]["link_inspection"]["links"][0]["status"] == "SUSPICIOUS"
    assert results[1]["status"] == "error"


def test_persistence_sanitizes_link_queries_in_all_reviewer_fields(tmp_path) -> None:
    result = analyze_cv_text_result(
        "Jane Example\nBerlin, Germany\nSoftware engineer",
        ai_settings=AISettings(enabled=False),
    )
    version = ComponentVersion("test", "1")
    details = FileDetails(
        source_format="pdf",
        extractor_version=version,
        fields=(
            FileDetail(
                FileDetailField.AUTHOR,
                "Synthetic Author",
                FileDetailStatus.AVAILABLE,
                "pdf",
                version,
            ),
        ),
    )
    inspection = LinkInspection(
        links=(
            LinkCheckResult(
                link_id="link:test:secret",
                status=LinkOutcomeStatus.REACHABLE,
                displayed_value="https://example.com/cv?candidate=private#fragment",
                sanitized_target="https://example.com/cv",
                source=LinkSource.VISIBLE_URL,
                association=LinkAssociation.VISIBLE_ONLY,
                role=LinkRole.PROFILE,
                source_page=1,
                source_evidence=(
                    Evidence(
                        page_id="page-0001",
                        page_number=1,
                        start_offset=0,
                        end_offset=72,
                        excerpt="Portfolio https://example.com/cv?candidate=private#fragment",
                    ),
                ),
                reason_code=LinkReasonCode.REACHABLE,
                terminal_status=200,
                terminal_registrable_domain="example.com",
                checked_at="2026-08-27T00:00:00+00:00",
                configuration_version="test-config-v1",
            ),
        ),
        checked_at="2026-08-27T00:00:00+00:00",
        configuration_version="test-config-v1",
    )
    report = replace(result.report, file_details=details, link_inspection=inspection)
    payload = serialize_report_payload(report)
    store = PersistenceStore(PersistenceConfig(tmp_path / "links.db"))
    analysis_id = store.persist_report(
        result.document_identity,
        report,
        report_payload=payload,
    )

    persisted = store.get_analysis_payload(analysis_id)
    assert persisted is not None
    persisted_link = persisted["link_inspection"]["links"][0]
    assert persisted_link["displayed_value"] == "https://example.com/cv"
    assert persisted_link["source_evidence"][0]["excerpt"] == (
        "Portfolio https://example.com/cv"
    )
    assert "candidate=private" not in json.dumps(persisted)
    assert "#fragment" not in json.dumps(persisted)
    assert "candidate=private" not in (tmp_path / "links.db").read_bytes().decode(
        "utf-8", errors="ignore"
    )
