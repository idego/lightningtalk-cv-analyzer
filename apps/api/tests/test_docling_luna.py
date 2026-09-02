from __future__ import annotations

from io import BytesIO
from threading import Lock
from time import perf_counter, sleep

import pytest
from docx import Document
from fastapi.testclient import TestClient
from reportlab.pdfgen.canvas import Canvas

from cv_validator.analysis.candidates import apply_review, validate_specialists
from cv_validator.analysis.docling_converter import DoclingTextConverter
from cv_validator.analysis.docling_luna import DoclingLunaAnalysisStrategy
from cv_validator.analysis.luna_client import (
    ModelPassError,
    ModelPassResponse,
    OpenAIResponsesLunaClient,
)
from cv_validator.analysis.source import SourceBlock, SourceDocument
from cv_validator.analysis.strategy import AnalysisInput, AnalysisStrategyError, SourceFormat
from cv_validator.api.app import create_app
from cv_validator.openai_config import OpenAISettings


def field(field_id: str, value: str, block_id: str, excerpt: str | None = None) -> dict:
    return {
        "id": field_id,
        "value": value,
        "evidence": [{"block_id": block_id, "excerpt": excerpt or value}],
    }


def pdf_bytes(*lines: str) -> bytes:
    output = BytesIO()
    canvas = Canvas(output)
    y = 760
    for line in lines:
        canvas.drawString(72, y, line)
        y -= 20
    canvas.save()
    return output.getvalue()


def docx_bytes() -> bytes:
    document = Document()
    document.add_heading("Alex Example", 0)
    document.add_paragraph("Developer in Opole, Poland")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Example University"
    table.cell(0, 1).text = "Computer Science"
    output = BytesIO()
    document.save(output)
    return output.getvalue()


class FakeLuna:
    def __init__(self, payloads: dict[str, dict], *, failing: set[str] | None = None, delay: float = 0) -> None:
        self.payloads = payloads
        self.failing = failing or set()
        self.delay = delay
        self.calls: list[tuple[str, float]] = []
        self.contexts: dict[str, dict | None] = {}
        self._lock = Lock()

    def run(self, pass_name, source, context=None):
        with self._lock:
            self.calls.append((pass_name, perf_counter()))
            self.contexts[pass_name] = context
        if self.delay and pass_name != "review":
            sleep(self.delay)
        if pass_name in self.failing:
            raise ModelPassError("fake_failure")
        return ModelPassResponse(
            self.payloads.get(pass_name, {}),
            "gpt-5.6-luna",
            {"input_tokens": 10, "output_tokens": 5, "total_tokens": 15},
        )


def complete_payloads() -> dict[str, dict]:
    return {
        "profile": {"profile": {
            "candidate_name": field("pf-name", "Alex Example", "texts/0"),
            "declared_location": field("pf-location", "Opole, Poland", "texts/1", "Developer in Opole, Poland"),
            "headline": field("pf-headline", "Developer", "texts/1", "Developer in Opole, Poland"),
            "summary": None,
            "skills": [field("pf-skill", "Python", "texts/2", "Python MongoDB")],
            "languages": [],
        }},
        "employment": {"records": [{
            "id": "employment-1",
            "organization": field("emp-org", "Example Systems", "texts/3", "Example Systems Developer"),
            "role": field("emp-role", "Developer", "texts/3", "Example Systems Developer"),
            "start_date": field("emp-start", "2022", "texts/4", "2022 - present"),
            "end_date": field("emp-end", "present", "texts/4", "2022 - present"),
            "location": None,
            "relationship_type": None,
        }, {
            "id": "employment-tech",
            "organization": field("tech-org", "MongoDB", "texts/2", "Python MongoDB"),
            "role": None,
            "start_date": None,
            "end_date": None,
            "location": None,
            "relationship_type": None,
        }]},
        "education": {"records": []},
        "review": {
            "accepted_record_ids": ["employment-1", "review-education-1"],
            "rejected_records": [{"id": "employment-tech", "reason_code": "technology_not_employer"}],
            "merge_groups": [],
            "relation_patches": [],
            "added_profile_fields": [],
            "added_candidates": [{
                "id": "review-education-1",
                "candidate_type": "education",
                "reason_code": "extractor_omission",
                "candidate": {
                    "institution": field("edu-inst", "Example University", "texts/5"),
                    "program": field("edu-program", "Computer Science", "texts/5", "Example University Computer Science"),
                    "degree": None,
                    "certificate": None,
                    "start_date": None,
                    "end_date": None,
                    "location": None,
                },
            }],
            "conflicts": [],
            "coverage_gaps": [],
            "status": "completed",
        },
    }


@pytest.mark.parametrize("source_format", [SourceFormat.PDF, SourceFormat.DOCX])
def test_docling_converts_text_pdf_and_docx_without_models(source_format) -> None:
    content = (
        pdf_bytes("Alex Example", "Developer in Opole, Poland")
        if source_format is SourceFormat.PDF
        else docx_bytes()
    )
    source = DoclingTextConverter().convert(content, f"candidate.{source_format.value}", source_format)

    assert source.blocks
    assert source.identity == DoclingTextConverter().convert(
        content, f"candidate.{source_format.value}", source_format
    ).identity
    assert [block.order for block in source.blocks] == list(range(len(source.blocks)))
    if source_format is SourceFormat.PDF:
        assert source.blocks[0].page_number == 1
        assert source.blocks[0].bbox is not None
    else:
        assert any(block.table_id and block.row_index == 0 for block in source.blocks)


def test_pdf_wrapped_lines_merge_into_one_block_but_entries_stay_separate() -> None:
    output = BytesIO()
    canvas = Canvas(output)
    canvas.setFont("Helvetica-Bold", 13)
    canvas.drawString(72, 760, "EDUCATION")
    canvas.setFont("Helvetica", 10)
    canvas.drawString(72, 740, "Master of Computer Science 2017 - 2020 Example University of Science and")
    canvas.drawString(72, 726, "Technology, Jilin, China")
    canvas.drawString(72, 712, "with distinction")
    canvas.drawString(72, 686, "Bachelor of Software Engineering 2013 - 2017 Example University,")
    canvas.drawString(72, 672, "Faisalabad, Punjab, Pakistan")
    canvas.drawString(90, 646, "- First bullet item")
    canvas.drawString(90, 632, "- Second bullet item")
    canvas.save()

    source = DoclingTextConverter().convert(output.getvalue(), "wrapped.pdf", SourceFormat.PDF)
    texts = [block.text for block in source.blocks]

    assert texts == [
        "EDUCATION",
        "Master of Computer Science 2017 - 2020 Example University of Science and Technology, Jilin, China with distinction",
        "Bachelor of Software Engineering 2013 - 2017 Example University, Faisalabad, Punjab, Pakistan",
        "- First bullet item",
        "- Second bullet item",
    ]
    assert [block.order for block in source.blocks] == list(range(len(source.blocks)))
    assert all(block.bbox is not None and block.page_number == 1 for block in source.blocks)


def test_scan_only_pdf_fails_with_clear_text_layer_error() -> None:
    output = BytesIO()
    canvas = Canvas(output)
    canvas.rect(50, 50, 100, 100, fill=1)
    canvas.save()

    with pytest.raises(AnalysisStrategyError, match="document_text_layer_unavailable"):
        DoclingTextConverter().convert(output.getvalue(), "scan.pdf", SourceFormat.PDF)


def test_far_fields_are_detached_and_records_stay_isolated() -> None:
    source = SourceDocument.create(tuple(
        SourceBlock(f"b-{index}", text, order=index)
        for index, text in enumerate([
            "Example Systems", "Developer", "2022", "filler", "filler", "filler",
            "Other Company", "Designer", "2024", "Example University",
        ])
    ), "pdf")
    employment = {"records": [{
        "id": "mixed",
        "organization": field("mixed-org", "Example Systems", "b-0"),
        "role": field("mixed-role", "Designer", "b-7"),
        "start_date": field("mixed-date", "2024", "b-8"),
        "end_date": None, "location": None, "relationship_type": None,
    }, {
        "id": "valid",
        "organization": field("valid-org", "Other Company", "b-6"),
        "role": field("valid-role", "Designer", "b-7"),
        "start_date": field("valid-date", "2024", "b-8"),
        "end_date": None, "location": None, "relationship_type": None,
    }]}
    state = validate_specialists(source, {}, employment, {})

    # Far fields are detached from the record instead of discarding it whole.
    mixed, valid = state.employment
    assert mixed["id"] == "mixed" and valid["id"] == "valid"
    assert mixed["organization"]["value"] == "Example Systems"
    assert mixed["role"] is None and mixed["start_date"] is None
    assert mixed["relation_status"] == "supported"
    assert valid["role"]["value"] == "Designer" and valid["start_date"]["value"] == "2024"
    detached = [c for c in state.conflicts if c["reason_code"] == "field_detached_from_record"]
    assert {(c["record_id"], c["field"]) for c in detached} == {("mixed", "role"), ("mixed", "start_date")}
    assert state.rejected == []


def test_fields_from_different_table_rows_are_detached_not_merged() -> None:
    source = SourceDocument.create((
        SourceBlock(
            "table/cell-0-0", "Example Systems", order=0,
            parent_id="table", table_id="table", row_index=0, column_index=0,
        ),
        SourceBlock(
            "table/cell-1-0", "Designer", order=1,
            parent_id="table", table_id="table", row_index=1, column_index=0,
        ),
    ), "docx")
    state = validate_specialists(source, {}, {"records": [{
        "id": "cross-row",
        "organization": field("cross-row-org", "Example Systems", "table/cell-0-0"),
        "role": field("cross-row-role", "Designer", "table/cell-1-0"),
        "start_date": None, "end_date": None, "location": None,
        "relationship_type": None,
    }]}, {})

    record, = state.employment
    assert record["organization"]["value"] == "Example Systems"
    assert record["role"] is None
    assert [c["field"] for c in state.conflicts if c["reason_code"] == "field_detached_from_record"] == ["role"]


def test_evidence_excerpt_must_contain_the_semantic_value() -> None:
    source = SourceDocument.create((
        SourceBlock("b-0", "Example Systems — Developer", order=0),
    ), "pdf")
    state = validate_specialists(source, {}, {"records": [{
        "id": "bad-excerpt",
        "organization": field("bad-org", "Example Systems", "b-0", "Developer"),
        "role": None, "start_date": None, "end_date": None,
        "location": None, "relationship_type": None,
    }]}, {})

    assert state.employment == []
    assert {item["reason_code"] for item in state.rejected} >= {"missing_organization"}


def test_value_spanning_adjacent_blocks_is_supported_by_ordered_excerpts() -> None:
    source = SourceDocument.create((
        SourceBlock("b-0", "Master of Computer Science 2017 - 2020 Changchun University of Science and", order=0),
        SourceBlock("b-1", "Technology, Jilin, China", order=1),
        SourceBlock("b-2", "filler", order=2),
        SourceBlock("b-3", "University of Agriculture,", order=3),
        SourceBlock("b-4", "Faisalabad, Punjab, Pakistan", order=4),
    ), "pdf")
    spanning = {
        "id": "inst-1",
        "value": "Changchun University of Science and Technology",
        "evidence": [
            {"block_id": "b-0", "excerpt": "Changchun University of Science and"},
            {"block_id": "b-1", "excerpt": "Technology"},
        ],
    }
    partial_second = {
        "id": "inst-2",
        "value": "University of Agriculture",
        "evidence": [
            {"block_id": "b-3", "excerpt": "University of Agriculture,"},
            {"block_id": "b-4", "excerpt": "Faisalabad, Punjab, Pakistan"},
        ],
    }
    non_adjacent = {
        "id": "inst-3",
        "value": "Changchun University of Science and Technology",
        "evidence": [
            {"block_id": "b-0", "excerpt": "Changchun University of Science and"},
            {"block_id": "b-4", "excerpt": "Faisalabad, Punjab, Pakistan"},
        ],
    }
    wrong_order = {
        "id": "inst-4",
        "value": "Changchun University of Science and Technology",
        "evidence": [
            {"block_id": "b-1", "excerpt": "Technology"},
            {"block_id": "b-0", "excerpt": "Changchun University of Science and"},
        ],
    }
    empty = {"institution": None, "program": None, "degree": None, "certificate": None,
             "start_date": None, "end_date": None, "location": None}
    state = validate_specialists(source, {}, {}, {"records": [
        {"id": "edu-1", **empty, "institution": spanning},
        {"id": "edu-2", **empty, "institution": partial_second},
        {"id": "edu-3", **empty, "institution": non_adjacent},
        {"id": "edu-4", **empty, "institution": wrong_order},
    ]})

    assert [record["id"] for record in state.education] == ["edu-1", "edu-2"]
    assert state.education[0]["institution"]["value"] == "Changchun University of Science and Technology"
    assert [item["source_id"] for item in state.education[0]["institution"]["evidence"]] == ["b-0", "b-1"]
    assert {item["id"] for item in state.rejected} == {"edu-3", "edu-4"}


def test_reviewer_unknown_ids_and_invalid_added_evidence_are_rejected() -> None:
    source = SourceDocument.create((SourceBlock("b-1", "Example University", order=0),), "pdf")
    state = validate_specialists(source, {}, {}, {})
    state, review = apply_review(source, state, {
        "accepted_record_ids": ["unknown"],
        "relation_patches": [{"record_id": "unknown", "field_ids": ["missing"]}],
        "added_candidates": [{
            "id": "review-education-1", "candidate_type": "education",
            "candidate": {
                "institution": field("added-inst", "Invented University", "b-1", "Invented University"),
                "program": None, "degree": None, "certificate": None,
                "start_date": None, "end_date": None, "location": None,
            },
        }],
        "coverage_gaps": [], "conflicts": [], "status": "partial",
    })

    reasons = {item["reason_code"] for item in review["conflicts"]}
    assert "unknown_reviewer_record_id" in reasons
    assert "unknown_reviewer_patch_id" in reasons
    assert "reviewer_added_candidate_invalid_evidence" in reasons
    assert not state.education
    assert review["coverage_gaps"]


def test_reviewer_applies_relation_patch_by_stable_field_ids() -> None:
    source = SourceDocument.create(tuple(
        SourceBlock(f"b-{index}", text, order=index)
        for index, text in enumerate([
            "Example Systems", "Developer", "filler", "filler", "filler", "Designer",
        ])
    ), "pdf")
    state = validate_specialists(source, {}, {"records": [{
        "id": "needs-correction",
        "organization": field("target-org", "Example Systems", "b-0"),
        "role": field("wrong-role", "Designer", "b-5"),
        "start_date": None, "end_date": None, "location": None, "relationship_type": None,
    }, {
        "id": "field-source",
        "organization": field("source-org", "Example Systems", "b-0"),
        "role": field("correct-role", "Developer", "b-1"),
        "start_date": None, "end_date": None, "location": None, "relationship_type": None,
    }]}, {})

    state, review = apply_review(source, state, {
        "accepted_record_ids": ["needs-correction"],
        "rejected_records": [{"id": "field-source", "reason_code": "field_source_only"}],
        "merge_groups": [],
        "relation_patches": [{"record_id": "needs-correction", "field_ids": ["correct-role"]}],
        "added_profile_fields": [], "added_candidates": [], "coverage_gaps": [],
        "conflicts": [], "status": "completed",
    })

    corrected = next(record for record in state.employment if record["id"] == "needs-correction")
    assert corrected["role"]["value"] == "Developer"
    assert corrected["relation_status"] == "supported"
    assert corrected["status"] == "accepted"
    assert review["relation_corrections"] == [{
        "record_id": "needs-correction", "field_ids": ["correct-role"],
    }]


def test_reviewer_merge_is_applied_deterministically() -> None:
    source = SourceDocument.create((
        SourceBlock("b-0", "Example Systems", order=0),
        SourceBlock("b-1", "Developer", order=1),
    ), "pdf")
    state = validate_specialists(source, {}, {"records": [{
        "id": "canonical",
        "organization": field("org-1", "Example Systems", "b-0"),
        "role": None, "start_date": None, "end_date": None,
        "location": None, "relationship_type": None,
    }, {
        "id": "duplicate",
        "organization": field("org-2", "Example Systems", "b-0"),
        "role": field("role-2", "Developer", "b-1"),
        "start_date": None, "end_date": None,
        "location": None, "relationship_type": None,
    }]}, {})

    state, review = apply_review(source, state, {
        "accepted_record_ids": ["duplicate"], "rejected_records": [],
        "merge_groups": [["canonical", "duplicate"]], "relation_patches": [],
        "added_profile_fields": [], "added_candidates": [], "coverage_gaps": [],
        "conflicts": [], "status": "completed",
    })

    assert [record["id"] for record in state.employment] == ["canonical"]
    assert state.employment[0]["role"]["value"] == "Developer"
    assert state.employment[0]["status"] == "accepted"
    assert review["merged_ids"] == [["canonical", "duplicate"]]


def test_reviewer_adds_supported_profile_and_employment_candidates() -> None:
    source = SourceDocument.create((
        SourceBlock("b-0", "Platform engineer focused on Python", order=0),
        SourceBlock("b-1", "Example Systems — Developer", order=1),
    ), "pdf")
    state = validate_specialists(source, {}, {}, {})

    state, review = apply_review(source, state, {
        "accepted_record_ids": [], "rejected_records": [], "merge_groups": [],
        "relation_patches": [],
        "added_profile_fields": [{
            "field_name": "summary",
            "field": field("review-summary", "Platform engineer", "b-0", "Platform engineer"),
        }],
        "added_candidates": [{
            "id": "review-employment-1", "candidate_type": "employment",
            "reason_code": "extractor_omission", "candidate": {
                "organization": field("review-org", "Example Systems", "b-1"),
                "role": field("review-role", "Developer", "b-1"),
                "start_date": None, "end_date": None, "location": None,
                "relationship_type": None,
            },
        }],
        "coverage_gaps": [], "conflicts": [], "status": "completed",
    })

    assert state.profile["summary"]["value"] == "Platform engineer"
    assert state.employment[0]["id"] == "review-employment-1"
    assert state.employment[0]["status"] == "accepted"
    assert review["added_profile_fields"] == ["summary"]
    assert review["added_candidate_ids"] == ["review-employment-1"]


def test_specialists_run_concurrently_and_reviewer_runs_after_them() -> None:
    content = pdf_bytes(
        "Alex Example", "Developer in Opole, Poland", "Python MongoDB",
        "Example Systems Developer", "2022 - present", "Example University Computer Science",
    )
    client = FakeLuna(complete_payloads(), delay=0.08)
    strategy = DoclingLunaAnalysisStrategy(client=client)
    request = AnalysisInput.from_upload(content, "candidate.pdf", "en")
    started = perf_counter()
    report = strategy.analyze(request)
    elapsed = perf_counter() - started

    calls = {name: timestamp for name, timestamp in client.calls}
    assert max(calls[name] for name in ("profile", "employment", "education")) - min(
        calls[name] for name in ("profile", "employment", "education")
    ) < 0.05
    assert calls["review"] >= max(calls[name] for name in ("profile", "employment", "education"))
    assert elapsed < 0.28
    assert report["strategy"]["name"] == "docling-luna"
    assert report["base_analysis"]["review"]["added_candidate_ids"] == ["review-education-1"]
    assert [item["id"] for item in report["base_analysis"]["education"]] == ["review-education-1"]
    assert "employment-tech" not in report["base_analysis"]["review"]["accepted_ids"]
    assert report["base_analysis"]["pass_statuses"]["profile"]["reasoning_effort"] == "none"
    assert report["base_analysis"]["pass_statuses"]["review"]["reasoning_effort"] == "low"
    assert client.contexts["review"]["mechanical"] is not None


def test_failed_specialist_does_not_remove_other_passes() -> None:
    payloads = complete_payloads()
    payloads["review"]["accepted_record_ids"] = ["review-education-1"]
    client = FakeLuna(payloads, failing={"employment"})
    strategy = DoclingLunaAnalysisStrategy(client=client)
    report = strategy.analyze(AnalysisInput.from_upload(
        pdf_bytes("Alex Example", "Developer in Opole, Poland", "Python MongoDB", "Example Systems Developer", "2022 - present", "Example University Computer Science"),
        "candidate.pdf", "en",
    ))

    assert report["base_analysis"]["status"] == "partial"
    assert report["base_analysis"]["profile"]["candidate_name"]["value"] == "Alex Example"
    assert report["base_analysis"]["education"][0]["added_by_reviewer"] is True
    assert report["base_analysis"]["pass_statuses"]["employment"]["attempt_count"] == 2
    assert [name for name, _ in client.calls].count("employment") == 2
    assert [name for name, _ in client.calls].count("profile") == 1
    assert [name for name, _ in client.calls].count("education") == 1


def test_upload_persistence_ownership_and_health_with_real_strategy(tmp_path) -> None:
    client_impl = FakeLuna(complete_payloads())
    strategy = DoclingLunaAnalysisStrategy(client=client_impl)
    client = TestClient(create_app(
        db_path=tmp_path / "docling-luna.db",
        openai_settings=OpenAISettings(enabled=False),
        analysis_strategy=strategy,
    ))
    owner = {"X-Analysis-Access-Token": "owner"}
    created = client.post(
        "/analyze",
        files={"file": ("candidate.pdf", pdf_bytes(
            "Alex Example", "Developer in Opole, Poland", "Python MongoDB",
            "Example Systems Developer", "2022 - present", "Example University Computer Science",
        ), "application/pdf")},
        headers=owner,
    )

    assert client.get("/health").json()["ready"] is True
    assert created.status_code == 200
    report = created.json()
    assert report["contract_version"] == "base-analysis-v2"
    assert client.get(f"/analyses/{report['analysis_id']}", headers=owner).status_code == 200
    assert client.get(f"/analyses/{report['analysis_id']}", headers={"X-Analysis-Access-Token": "other"}).status_code == 404


def test_openai_contract_pins_model_store_and_reasoning() -> None:
    class Usage:
        def model_dump(self):
            return {"input_tokens": 1, "output_tokens": 1, "total_tokens": 2}

    class Response:
        output_text = "{}"
        model = "gpt-5.6-luna"
        usage = Usage()

    class Responses:
        def __init__(self):
            self.calls = []

        def create(self, **kwargs):
            self.calls.append(kwargs)
            return Response()

    class Client:
        def __init__(self):
            self.responses = Responses()

    raw_client = Client()
    client = OpenAIResponsesLunaClient(client=raw_client)
    source = SourceDocument.create((SourceBlock("b-1", "Alex Example", order=0),), "pdf")

    client.run("profile", source)
    client.run("review", source, {})

    specialist, reviewer = raw_client.responses.calls
    assert specialist["model"] == reviewer["model"] == "gpt-5.6-luna"
    assert specialist["reasoning"] == {"effort": "none"}
    assert reviewer["reasoning"] == {"effort": "low"}
    assert specialist["store"] is reviewer["store"] is False
    assert "tools" not in specialist and "tools" not in reviewer
    assert specialist["text"]["format"]["strict"] is True


def test_validated_records_are_accepted_by_default_and_reviewer_can_reject() -> None:
    source = SourceDocument.create((
        SourceBlock("b-0", "Example Systems — Developer", order=0),
        SourceBlock("b-1", "Served 100+ clients", order=1),
    ), "pdf")
    empty = {"role": None, "start_date": None, "end_date": None, "location": None, "relationship_type": None}

    def records() -> dict:
        return {"records": [
            {"id": "employment_1", **empty, "organization": field("org-1", "Example Systems", "b-0")},
            {"id": "employment_2", **empty, "organization": field("org-2", "100+ clients", "b-1")},
        ]}

    without_review, review = apply_review(source, validate_specialists(source, {}, records(), {}), {})
    assert [record["status"] for record in without_review.employment] == ["accepted", "accepted"]
    assert review["accepted_ids"] == ["employment_1", "employment_2"]

    reviewed, review = apply_review(source, validate_specialists(source, {}, records(), {}), {
        "accepted_record_ids": [],
        "rejected_records": [{"id": "employment_2", "reason_code": "client_count_not_employer"}],
        "merge_groups": [], "relation_patches": [], "added_profile_fields": [],
        "added_candidates": [], "conflicts": [], "coverage_gaps": [], "status": "completed",
    })
    assert {record["id"]: record["status"] for record in reviewed.employment} == {
        "employment_1": "accepted",
        "employment_2": "ambiguous",
    }
    assert review["accepted_ids"] == ["employment_1"]
    assert {"id": "employment_2", "reason_code": "client_count_not_employer"} in review["rejected"]


def test_certificate_only_education_record_is_kept() -> None:
    source = SourceDocument.create((
        SourceBlock("b-0", "AWS Certified Cloud Practitioner (CLF-C02) - In Progress", order=0),
        SourceBlock("b-1", "Computer Science", order=1),
    ), "pdf")
    empty = {"institution": None, "program": None, "degree": None, "certificate": None,
             "start_date": None, "end_date": None, "location": None}
    state = validate_specialists(source, {}, {}, {"records": [
        {"id": "education_1", **empty, "certificate": field(
            "cert", "AWS Certified Cloud Practitioner", "b-0", "AWS Certified Cloud Practitioner (CLF-C02)",
        )},
        {"id": "education_2", **empty, "program": field("prog", "Computer Science", "b-1")},
    ]})

    assert [record["id"] for record in state.education] == ["education_1"]
    assert {"id": "education_2", "reason_code": "missing_institution_or_certificate"} in state.rejected


def test_review_context_exposes_field_ids_and_bounds_reviewer_text() -> None:
    source = SourceDocument.create((
        SourceBlock("b-0", "Example Systems — Developer", order=0),
    ), "pdf")
    state = validate_specialists(source, {}, {"records": [{
        "id": "employment_1",
        "organization": field("employment_1.organization", "Example Systems", "b-0"),
        "role": None, "start_date": None, "end_date": None, "location": None, "relationship_type": None,
    }]}, {})

    from cv_validator.analysis.candidates import EMPLOYMENT_FIELDS, public_records
    context = public_records(state.employment, EMPLOYMENT_FIELDS, include_field_ids=True)
    assert context[0]["organization"]["field_id"] == "employment_1.organization"
    assert "field_id" not in public_records(state.employment, EMPLOYMENT_FIELDS)[0]["organization"]

    _, review = apply_review(source, state, {
        "accepted_record_ids": [], "rejected_records": [], "merge_groups": [],
        "relation_patches": [], "added_profile_fields": [], "added_candidates": [],
        "conflicts": [{
            "reason_code": "x" * 200, "record_ids": ["employment_1"], "field_ids": [],
            "source_block_ids": ["b-0"], "summary": "quoted cv text " * 40,
        }],
        "coverage_gaps": [{"target": "profile", "reason_code": "", "source_block_ids": []}],
        "status": "partial",
    })
    conflict = review["conflicts"][-1]
    assert conflict["reason_code"] == "reviewer_annotation"
    assert conflict["summary"] is None
    assert review["coverage_gaps"][0]["reason_code"] == "reviewer_annotation"


def test_docx_inline_formatting_runs_form_one_block() -> None:
    document = Document()
    paragraph = document.add_paragraph("Backend engineer with ")
    paragraph.add_run("Python").bold = True
    paragraph.add_run(", ")
    paragraph.add_run("Django").bold = True
    paragraph.add_run(" and cloud experience.")
    document.add_paragraph("Second paragraph")
    output = BytesIO()
    document.save(output)

    source = DoclingTextConverter().convert(output.getvalue(), "runs.docx", SourceFormat.DOCX)
    texts = [block.text for block in source.blocks]

    assert "Backend engineer with Python, Django and cloud experience." in texts
    assert "Second paragraph" in texts
    assert [block.order for block in source.blocks] == list(range(len(source.blocks)))


def test_truncated_model_output_is_reported_as_truncated() -> None:
    class Response:
        status = "incomplete"
        output_text = '{"profile": {'
        usage = None
        model = "gpt-5.6-luna"

    class Responses:
        def create(self, **kwargs):
            assert kwargs["max_output_tokens"] >= 6000
            return Response()

    class Client:
        responses = Responses()

    source = SourceDocument.create((SourceBlock("b-0", "Alex Example", order=0),), "pdf")
    with pytest.raises(ModelPassError, match="truncated"):
        OpenAIResponsesLunaClient(client=Client()).run("profile", source)


def test_company_research_rejections_carry_rule_names() -> None:
    from cv_validator.research.company import CompanyResearchRequest, validate_company_research
    from cv_validator.research.domain import CompanyResearchInvalidResponse

    request = CompanyResearchRequest(({"organization": "Example Systems"},))
    organization = {
        "query_subject": "Example Systems",
        "existence": "supported",
        "activity": None, "operating_dates": None, "location": None, "official_website": None,
        "company_pages": [], "registries": [], "findings": [{"claim": "x", "source_url": "https://example.com"}],
        "confidence": "low", "uncertainty": "n/a",
        "limited_online_presence": True,
        "limited_online_presence_reason": "does not establish existence or absence",
    }
    payload = {"organizations": [organization], "searches_performed": ["q"], "search_limitations": ["l"]}
    with pytest.raises(CompanyResearchInvalidResponse) as info:
        validate_company_research(payload, request=request)
    assert info.value.reason in {"schema", "limited_presence_contradiction"}
