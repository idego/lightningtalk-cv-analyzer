from __future__ import annotations

import io

from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.shared import Pt
from reportlab.pdfgen import canvas

from cv_validator.ai.config import AISettings
from cv_validator.pipeline import analyze_cv_bytes_result


def _docx(*, hidden: bool = False, tiny: bool = False, header: bool = False) -> bytes:
    stream = io.BytesIO()
    document = Document()
    if header:
        document.sections[0].header.paragraphs[0].text = "Unsupported header content"
    document.add_paragraph("Experience")
    document.add_paragraph("Acme Engineer 01/2020 - 03/2020")
    paragraph = document.add_paragraph()
    run = paragraph.add_run("Hidden review marker")
    run.font.hidden = hidden
    if tiny:
        run.font.size = Pt(0.5)
    document.add_paragraph("Berlin Germany phone +49 30 123456")
    document.save(stream)
    return stream.getvalue()


def _pdf(*, tiny: bool = False, white_on_light: bool = False) -> bytes:
    stream = io.BytesIO()
    document = canvas.Canvas(stream)
    document.setFont("Helvetica", 11)
    document.drawString(72, 760, "Experience")
    document.drawString(72, 740, "Acme Engineer 01/2020 - 03/2020")
    if white_on_light:
        document.setFillColorRGB(0.97, 0.97, 0.97)
        document.rect(68, 712, 180, 18, fill=1, stroke=0)
        document.setFillColorRGB(1, 1, 1)
    document.setFont("Helvetica", 0.5 if tiny else 11)
    document.drawString(72, 720, "Tiny review marker")
    document.setFont("Helvetica", 11)
    document.drawString(72, 700, "Berlin Germany phone +49 30 123456")
    document.save()
    return stream.getvalue()


def test_hidden_docx_and_unsupported_header_return_safe_partial_result():
    result = analyze_cv_bytes_result(_docx(hidden=True, header=True), "cv.docx", ai_settings=AISettings(enabled=False))
    audits = result.report.structural_audits
    assert audits.status.value == "partial"
    assert "docx_headers" in audits.coverage.omitted_parts
    assert audits.visibility.observations[0].kind == "hidden_text"
    assert "Hidden review marker" not in str(audits.visibility.to_dict() if hasattr(audits.visibility, "to_dict") else audits.visibility)


def test_tiny_docx_is_detected_without_failing_base_analysis():
    audits = analyze_cv_bytes_result(_docx(tiny=True), "cv.docx", ai_settings=AISettings(enabled=False)).report.structural_audits
    assert any(item.kind == "near_zero_text" for item in audits.visibility.observations)


def test_tiny_pdf_is_detected_from_bounded_character_provenance():
    audits = analyze_cv_bytes_result(_pdf(tiny=True), "cv.pdf", ai_settings=AISettings(enabled=False)).report.structural_audits
    assert any(item.kind == "near_zero_text" for item in audits.visibility.observations)
    assert audits.coverage.audited_parts == ("pdf_page_text_spans",)


def test_ordinary_pdf_has_no_visibility_finding():
    audits = analyze_cv_bytes_result(_pdf(), "cv.pdf", ai_settings=AISettings(enabled=False)).report.structural_audits
    assert audits.visibility.observations == ()


def test_white_pdf_text_requires_an_explicit_known_light_background():
    audits = analyze_cv_bytes_result(_pdf(white_on_light=True), "cv.pdf", ai_settings=AISettings(enabled=False)).report.structural_audits
    assert any(item.kind == "low_contrast_text" for item in audits.visibility.observations)
