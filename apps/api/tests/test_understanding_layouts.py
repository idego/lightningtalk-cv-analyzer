from io import BytesIO

import pytest
from docx import Document

from cv_validator.document_understanding.service import understand_document, understanding_to_payload
from cv_validator.ingestion import PresentationSpan, RawDocument, SourcePage
from cv_validator.ingestion.docx import extract_docx
from cv_validator.ingestion.redaction import redact_national_ids


def test_docx_table_rows_are_independent_entry_boundaries():
    document = Document(); document.add_paragraph("Experience")
    table = document.add_table(rows=2, cols=3)
    values = (("First Company Ltd", "Software Engineer", "01/2020 - 02/2022"), ("Second Company Ltd", "Product Manager", "03/2022 - Present"))
    for row, row_values in zip(table.rows, values):
        for cell, value in zip(row.cells, row_values): cell.text = value
    buffer = BytesIO(); document.save(buffer)
    parsed = extract_docx(buffer.getvalue())
    payload = understanding_to_payload(understand_document(redact_national_ids(parsed), "test", snapshot_month="2026-08"))
    assert [next(field["value"] for field in record["fields"] if field["name"] == "organization") for record in payload["records"]] == ["First Company Ltd", "Second Company Ltd"]


def test_nested_docx_table_preserves_source_path_and_emphasis_for_employment():
    document = Document(); document.add_paragraph("Experience")
    outer = document.add_table(rows=1, cols=1)
    nested = outer.cell(0, 0).add_table(rows=1, cols=2)
    identity = nested.cell(0, 0).paragraphs[0]
    role = identity.add_run("Platform Steward"); role.bold = True
    identity.add_run(" | Example Orbit Ltd")
    nested.cell(0, 1).text = "Jan 2020 - Feb 2022"
    buffer = BytesIO(); document.save(buffer)

    parsed = extract_docx(buffer.getvalue())
    role_span = next(span for span in parsed.presentation_spans if span.text == "Platform Steward")
    source_block = next(block for block in parsed.source_blocks if role_span.start_offset in range(block.start_offset or -1, block.end_offset or -1))
    assert role_span.bold is True
    assert role_span.paragraph_path.rsplit("/run/", 1)[0] == source_block.paragraph_path

    payload = understanding_to_payload(understand_document(redact_national_ids(parsed), "test", snapshot_month="2026-08"))
    employment = next(record for record in payload["records"] if record["kind"] == "employment")
    fields = {field["name"]: field["value"] for field in employment["fields"]}
    assert fields["role"] == "Platform Steward"
    assert fields["organization"] == "Example Orbit Ltd"
    assert fields["employment_dates"] == "Jan 2020 - Feb 2022"


def test_emphasized_role_pairs_with_one_date_anchored_unlabelled_employer():
    document = Document(); document.add_paragraph("Experience")
    identity = document.add_paragraph()
    role = identity.add_run("Systems Engineer"); role.bold = True
    identity.add_run(" | Northwind Workshop")
    document.add_paragraph("Jan 2021 - Mar 2024")
    buffer = BytesIO(); document.save(buffer)

    payload = understanding_to_payload(understand_document(redact_national_ids(extract_docx(buffer.getvalue())), "test", snapshot_month="2026-08"))
    employment = next(record for record in payload["records"] if record["kind"] == "employment")
    fields = {field["name"]: field["value"] for field in employment["fields"]}
    assert fields["role"] == "Systems Engineer"
    assert fields["organization"] == "Northwind Workshop"


@pytest.mark.parametrize("separator", [" at ", " — "])
def test_bounded_same_line_connector_pairs_role_and_organization(separator):
    text = f"Experience\nSoftware Engineer{separator}Example Workshop — Jan 2021 - Mar 2024"
    payload = understanding_to_payload(understand_document(redact_national_ids(RawDocument(pages=(SourcePage("page-0001", 1, text),), source_format="text")), "test", snapshot_month="2026-08"))
    employment = next(record for record in payload["records"] if record["kind"] == "employment")
    fields = {field["name"]: field["value"] for field in employment["fields"]}
    assert fields["role"] == "Software Engineer"
    assert fields["organization"] == "Example Workshop"


@pytest.mark.parametrize("other", ["Remote", "Taylor Morgan", "Built APIs"])
def test_generic_short_fragment_cannot_establish_employer_ownership(other):
    text = f"Experience\nSoftware Engineer\n{other}\nJan 2021 - Mar 2024"
    payload = understanding_to_payload(understand_document(redact_national_ids(RawDocument(pages=(SourcePage("page-0001", 1, text),), source_format="text")), "test", snapshot_month="2026-08"))
    assert [record for record in payload["records"] if record["kind"] == "employment"] == []
    assert payload["code_research_subjects"] == []


def test_uniform_font_size_does_not_turn_duty_text_into_role():
    text = "Experience\nCompany: Example Labs Ltd\nBuilt APIs\nJan 2021 - Mar 2024"
    page = SourcePage("page-0001", 1, text)
    spans = tuple(PresentationSpan(page.page_id, 1, line.text, line.start_offset, line.end_offset, association="exact", font_size_points=11) for line in page.lines)
    raw = RawDocument(pages=(page,), source_format="pdf", presentation_spans=spans)
    payload = understanding_to_payload(understand_document(redact_national_ids(raw), "test", snapshot_month="2026-08"))
    record = next(record for record in payload["records"] if record["kind"] == "employment")
    role = next(field for field in record["fields"] if field["name"] == "role")
    assert role["status"] == "unknown"
    assert role["value"] is None


def test_styled_identity_without_date_anchor_does_not_create_employment_record():
    document = Document(); document.add_paragraph("Experience")
    identity = document.add_paragraph()
    role = identity.add_run("Systems Engineer"); role.bold = True
    identity.add_run(" | Northwind Workshop")
    document.add_paragraph("Built internal tools and improved operations.")
    buffer = BytesIO(); document.save(buffer)

    payload = understanding_to_payload(understand_document(redact_national_ids(extract_docx(buffer.getvalue())), "test", snapshot_month="2026-08"))
    assert [record for record in payload["records"] if record["kind"] == "employment"] == []


def test_interleaved_pdf_column_ownership_abstains_instead_of_guessing():
    text = "Experience\nFirst Company Ltd\nSecond Company Ltd\nSoftware Engineer\n01/2020 - 02/2022"
    spans = tuple(PresentationSpan("page-0001", 1, line.text, line.start_offset, line.end_offset, association="exact", bbox=(0, index * 10, 100, index * 10 + 8)) for index, line in enumerate(SourcePage("page-0001", 1, text).lines))
    raw = RawDocument(pages=(SourcePage("page-0001", 1, text),), source_format="pdf", presentation_spans=spans)
    payload = understanding_to_payload(understand_document(redact_national_ids(raw), "test", snapshot_month="2026-08"))
    assert payload["records"] == []
    assert any(item["reason_code"] == "unsupported_employment_identity" for item in payload["ambiguous_spans"])


def test_oversized_employment_entry_abstains_with_explicit_limit_issue():
    lines = ["Experience", "Jan 2021 - Mar 2024", *[f"Responsibility item {index}" for index in range(80)]]
    text = "\n".join(lines)
    payload = understanding_to_payload(understand_document(redact_national_ids(RawDocument(pages=(SourcePage("page-0001", 1, text),), source_format="text")), "test", snapshot_month="2026-08"))
    assert [record for record in payload["records"] if record["kind"] == "employment"] == []
    assert any(item["reason_code"] == "employment_candidate_limit" for item in payload["ambiguous_spans"])


def test_duties_between_date_anchored_jobs_do_not_own_the_next_identity():
    text = """Experience
Software Engineer | First Company Ltd
Jan 2020 - Feb 2022
• Led platform migration
• Managed release support
Product Manager | Second Company Ltd
Mar 2022 - Present"""
    payload = understanding_to_payload(understand_document(redact_national_ids(RawDocument(pages=(SourcePage("page-0001", 1, text),), source_format="text")), "test", snapshot_month="2026-08"))
    records = [record for record in payload["records"] if record["kind"] == "employment"]
    fields = [{field["name"]: field["value"] for field in record["fields"]} for record in records]

    assert [(item["role"], item["organization"], item["employment_dates"]) for item in fields] == [
        ("Software Engineer", "First Company Ltd", "Jan 2020 - Feb 2022"),
        ("Product Manager", "Second Company Ltd", "Mar 2022 - Present"),
    ]
    assert all("Led" not in (item["role"] or "") and "Managed" not in (item["role"] or "") for item in fields)
    assert len(payload["timeline_record_links"]) == 2
    assert len({link["record_id"] for link in payload["timeline_record_links"]}) == 2


def test_docx_list_item_duty_does_not_become_employment_identity():
    document = Document(); document.add_paragraph("Experience")
    document.add_paragraph("Software Engineer | First Company Ltd")
    document.add_paragraph("Jan 2020 - Feb 2022")
    document.add_paragraph("Led platform migration", style="List Bullet")
    document.add_paragraph("Product Manager | Second Company Ltd")
    document.add_paragraph("Mar 2022 - Present")
    buffer = BytesIO(); document.save(buffer)

    payload = understanding_to_payload(understand_document(redact_national_ids(extract_docx(buffer.getvalue())), "test", snapshot_month="2026-08"))
    records = [record for record in payload["records"] if record["kind"] == "employment"]
    fields = [{field["name"]: field["value"] for field in record["fields"]} for record in records]
    assert [(item["role"], item["organization"]) for item in fields] == [
        ("Software Engineer", "First Company Ltd"),
        ("Product Manager", "Second Company Ltd"),
    ]


def test_combined_education_line_uses_bounded_semantic_fragments():
    text = "Education\nExample University | MSc Computer Science | 2020 - 2022"
    payload = understanding_to_payload(understand_document(redact_national_ids(RawDocument(pages=(SourcePage("page-0001", 1, text),), source_format="text")), "test", snapshot_month="2026-08"))
    record = next(record for record in payload["records"] if record["kind"] == "education")
    fields = {field["name"]: field for field in record["fields"]}

    assert fields["institution"]["value"] == "Example University"
    assert fields["degree"]["value"] == "MSc Computer Science"
    assert fields["study_dates"]["value"] == "2020 - 2022"
    assert fields["institution"]["evidence"][0]["excerpt"] == "Example University"


def test_unbounded_combined_education_line_abstains_instead_of_storing_whole_line():
    text = "Education\nExample University MSc Computer Science 2020 - 2022"
    payload = understanding_to_payload(understand_document(redact_national_ids(RawDocument(pages=(SourcePage("page-0001", 1, text),), source_format="text")), "test", snapshot_month="2026-08"))
    assert [record for record in payload["records"] if record["kind"] == "education"] == []
