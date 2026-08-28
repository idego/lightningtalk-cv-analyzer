import io
from pathlib import Path

import pytest
from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from reportlab.pdfgen import canvas

from cv_validator.config import IngestionConfig
from cv_validator.domain import Band
from cv_validator.ingestion import (
    EmptyTextError,
    IngestionError,
    InsufficientTextError,
    RawDocument,
    SourcePage,
)
from cv_validator.ingestion.docx import extract_docx
from cv_validator.ingestion.pdf import extract_pdf
from cv_validator.ingestion.redaction import redact_national_ids
from cv_validator.ingestion.router import ingest_cv
from cv_validator.ingestion.text import (
    meaningful_token_count,
    validate_text_sufficiency,
)
from cv_validator.pipeline import analyze_cv_text

FIXTURES = Path(__file__).parent.parent / "fixtures" / "calibration"


def _make_text_pdf(text: str) -> bytes:
    buffer = io.BytesIO()
    c = canvas.Canvas(buffer)
    y = 800
    for line in text.splitlines():
        c.drawString(72, y, line)
        y -= 14
    c.save()
    return buffer.getvalue()


def _make_page_separated_pdf(*page_texts: str) -> bytes:
    buffer = io.BytesIO()
    c = canvas.Canvas(buffer)
    for page_index, text in enumerate(page_texts):
        y = 800
        for line in text.splitlines():
            c.drawString(72, y, line)
            y -= 14
        if page_index < len(page_texts) - 1:
            c.showPage()
    c.save()
    return buffer.getvalue()


def _make_scanned_pdf() -> bytes:
    buffer = io.BytesIO()
    c = canvas.Canvas(buffer)
    c.save()
    return buffer.getvalue()


def _make_docx(text: str) -> bytes:
    doc = Document()
    for line in text.splitlines():
        doc.add_paragraph(line)
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _save_docx(doc: Document) -> bytes:
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


SAMPLE = """Jane Doe
Berlin, Germany
+49 30 12345678

Experience
Engineer — Acme, Berlin
"""


def test_pdf_text_extraction():
    content = _make_text_pdf(SAMPLE)
    parsed = extract_pdf(content)
    assert "Berlin, Germany" in parsed.pages[0].text
    assert parsed.source_lines
    assert parsed.source_blocks
    assert parsed.source_blocks[0].bbox is not None


def test_pdf_preserves_two_real_pages_and_stable_ids():
    parsed = extract_pdf(
        _make_page_separated_pdf(
            "Alpha Bravo Charlie",
            "Delta Echo Foxtrot",
        )
    )

    assert [page.page_id for page in parsed.pages] == ["page-0001", "page-0002"]
    assert [page.page_number for page in parsed.pages] == [1, 2]
    assert [page.text for page in parsed.pages] == [
        "Alpha Bravo Charlie",
        "Delta Echo Foxtrot",
    ]


def test_docx_extraction():
    parsed = extract_docx(_make_docx(SAMPLE))
    assert "Berlin, Germany" in parsed.pages[0].text


def test_docx_without_explicit_page_break_is_one_page():
    parsed = extract_docx(_make_docx("Alpha Bravo Charlie\nDelta Echo Foxtrot"))

    assert len(parsed.pages) == 1
    assert parsed.pages[0].page_id == "page-0001"


def test_docx_hard_page_break_creates_logical_page():
    doc = Document()
    paragraph = doc.add_paragraph()
    paragraph.add_run("Alpha Bravo Charlie")
    paragraph.add_run().add_break(WD_BREAK.PAGE)
    paragraph.add_run("Delta Echo Foxtrot")

    parsed = extract_docx(_save_docx(doc))

    assert [page.page_id for page in parsed.pages] == ["page-0001", "page-0002"]
    assert parsed.pages[0].text == "Alpha Bravo Charlie"
    assert parsed.pages[1].text == "Delta Echo Foxtrot"


def test_docx_page_break_before_creates_logical_page():
    doc = Document()
    doc.add_paragraph("Alpha Bravo Charlie")
    second_page = doc.add_paragraph("Delta Echo Foxtrot")
    second_page.paragraph_format.page_break_before = True

    parsed = extract_docx(_save_docx(doc))

    assert [page.text for page in parsed.pages] == [
        "Alpha Bravo Charlie",
        "Delta Echo Foxtrot",
    ]


def test_docx_flattens_table_text_in_document_block_order():
    doc = Document()
    doc.add_paragraph("Profile begins here")
    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Jane Doe"
    table.cell(0, 1).text = "Berlin Germany"
    doc.add_paragraph("Experience Software Engineer")

    parsed = extract_docx(_save_docx(doc))

    assert parsed.pages[0].text.splitlines() == [
        "Profile begins here",
        "Jane Doe",
        "Berlin Germany",
        "Experience Software Engineer",
    ]
    table_blocks = [block for block in parsed.source_blocks if block.kind == "table_cell"]
    assert [block.row_index for block in table_blocks] == [0, 0]
    assert len({block.table_id for block in table_blocks}) == 1
    assert all(block.association == "exact" for block in table_blocks)


def test_docx_hard_page_break_inside_table_is_preserved():
    doc = Document()
    cell_paragraph = doc.add_table(rows=1, cols=1).cell(0, 0).paragraphs[0]
    cell_paragraph.add_run("Alpha Bravo Charlie")
    cell_paragraph.add_run().add_break(WD_BREAK.PAGE)
    cell_paragraph.add_run("Delta Echo Foxtrot")

    parsed = extract_docx(_save_docx(doc))

    assert [page.text for page in parsed.pages] == [
        "Alpha Bravo Charlie",
        "Delta Echo Foxtrot",
    ]


def test_docx_page_break_before_inside_table_is_preserved():
    doc = Document()
    doc.add_paragraph("Alpha Bravo Charlie")
    cell_paragraph = doc.add_table(rows=1, cols=1).cell(0, 0).paragraphs[0]
    cell_paragraph.add_run("Delta Echo Foxtrot")
    cell_paragraph.paragraph_format.page_break_before = True

    parsed = extract_docx(_save_docx(doc))

    assert [page.text for page in parsed.pages] == [
        "Alpha Bravo Charlie",
        "Delta Echo Foxtrot",
    ]


def test_docx_ignores_last_rendered_page_break_inside_table():
    doc = Document()
    paragraph = doc.add_table(rows=1, cols=1).cell(0, 0).paragraphs[0]
    first_run = paragraph.add_run("Alpha Bravo Charlie")
    first_run._r.append(OxmlElement("w:lastRenderedPageBreak"))
    paragraph.add_run(" Delta Echo Foxtrot")

    parsed = extract_docx(_save_docx(doc))

    assert len(parsed.pages) == 1
    assert parsed.pages[0].text == "Alpha Bravo Charlie Delta Echo Foxtrot"


def test_docx_preserves_source_text_without_layout_reconstruction():
    doc = Document()
    paragraph = doc.add_paragraph()
    paragraph.add_run("  Alpha").add_tab()
    paragraph.add_run("Bravo  ").add_break(WD_BREAK.LINE)
    paragraph.add_run("Charlie")
    doc.add_paragraph("")
    doc.add_paragraph("Delta Echo")

    parsed = extract_docx(_save_docx(doc))

    assert parsed.pages[0].text == "  Alpha\tBravo  \nCharlie\n\nDelta Echo"


def test_source_lines_have_page_local_numbers_and_exact_offsets():
    parsed = RawDocument(
        pages=(
            SourcePage("page-0001", 1, "Alpha\nBravo Charlie"),
            SourcePage("page-0002", 2, "Delta\nEcho Foxtrot"),
        ),
        source_format="text",
    )

    assert [line.line_number for line in parsed.pages[0].lines] == [1, 2]
    assert [line.line_number for line in parsed.pages[1].lines] == [1, 2]
    assert [
        (line.start_offset, line.end_offset) for line in parsed.pages[0].lines
    ] == [(0, 5), (6, 19)]
    for page in parsed.pages:
        for line in page.lines:
            assert page.text[line.start_offset : line.end_offset] == line.text


def test_page_separated_markdown_preserves_exact_source_text():
    raw = RawDocument(
        pages=(
            SourcePage("page-0001", 1, "Alpha\nBravo"),
            SourcePage("page-0002", 2, "Charlie\nDelta"),
        ),
        source_format="text",
    )
    parsed = redact_national_ids(raw)

    assert parsed.markdown == (
        "<!-- page: page-0001 -->\nAlpha\nBravo\n\n"
        "<!-- page: page-0002 -->\nCharlie\nDelta"
    )


def test_zero_meaningful_tokens_is_empty_text():
    parsed = RawDocument(
        pages=(SourcePage("page-0001", 1, "... — x 7"),),
        source_format="text",
    )

    with pytest.raises(EmptyTextError):
        validate_text_sufficiency(parsed, IngestionConfig())


@pytest.mark.parametrize("token_count", range(1, 5))
def test_one_to_four_meaningful_tokens_are_insufficient(token_count):
    text = " ".join(["alpha", "bravo", "charlie", "delta"][:token_count])
    parsed = RawDocument(
        pages=(SourcePage("page-0001", 1, text),),
        source_format="text",
    )

    with pytest.raises(InsufficientTextError):
        validate_text_sufficiency(parsed, IngestionConfig())


def test_five_meaningful_tokens_are_accepted():
    parsed = RawDocument(
        pages=(SourcePage("page-0001", 1, "alpha bravo charlie delta echo"),),
        source_format="text",
    )

    validate_text_sufficiency(parsed, IngestionConfig())


def test_existing_sparse_cv_is_accepted_and_remains_gray():
    text = (FIXTURES / "sparse_cv.txt").read_text()

    parsed = extract_docx(_make_docx(text))
    report = analyze_cv_text(text)

    assert meaningful_token_count(parsed) >= 5
    assert report.band is Band.GRAY


def test_reject_scanned_pdf():
    with pytest.raises(IngestionError, match="no extractable text"):
        extract_pdf(_make_scanned_pdf())


def test_reject_unsupported_format():
    with pytest.raises(IngestionError, match="Unsupported"):
        ingest_cv(b"hello", filename="cv.txt")


def test_reject_empty_docx():
    with pytest.raises(IngestionError, match="no extractable text"):
        extract_docx(_make_docx(""))
