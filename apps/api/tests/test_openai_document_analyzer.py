import json
import io
from types import SimpleNamespace

import httpx
import openai
import pytest
from docx import Document
from fastapi.testclient import TestClient

from cv_validator.ai.application import (
    DocumentAnalyzerClientError,
    DocumentAnalyzerTimeoutError,
)
from cv_validator.ai.config import AISettings
from cv_validator.ai.domain import (
    AIAnalysisStatus,
    AIFailureReason,
    DocumentAnalyzerResponse,
)
from cv_validator.ai.openai_client import OpenAIResponsesDocumentAnalyzer
from cv_validator.ai.request import SCHEMA_VERSION, build_document_analysis_request
from cv_validator.api.app import create_app
from cv_validator.extraction.deterministic import analyze_deterministically
from cv_validator.ingestion import RawDocument, SourcePage
from cv_validator.ingestion.redaction import redact_national_ids
from cv_validator.pipeline import analyze_cv_text_result
from cv_validator.serialization import serialize_report_payload


def _empty_response() -> dict:
    return {
        "facts": {"contact": [], "education": [], "employment": []},
        "findings": [],
        "unknowns": [],
        "analysis_limitations": [],
    }


class _Usage:
    def model_dump(self):
        return {"input_tokens": 12, "output_tokens": 4}


class _Responses:
    def __init__(self, response=None, error=None):
        self.response = response
        self.error = error
        self.calls = []

    def create(self, **payload):
        self.calls.append(payload)
        if self.error is not None:
            raise self.error
        return self.response


class _Analyzer:
    def __init__(self, response=None, error=None):
        self.response = response
        self.error = error
        self.requests = []

    def analyze(self, request):
        self.requests.append(request)
        if self.error is not None:
            raise self.error
        return self.response


def _request():
    document = redact_national_ids(
        RawDocument(
            pages=(SourcePage("page-0001", 1, "Candidate\nEngineer profile"),),
            source_format="text",
        )
    )
    deterministic = analyze_deterministically(document, "1.0.0")
    settings = AISettings(enabled=True, api_key="test-key")
    return settings, build_document_analysis_request(
        settings,
        document,
        deterministic,
    )


def test_responses_analyzer_sends_the_exact_independent_production_payload() -> None:
    settings, request = _request()
    responses = _Responses(
        SimpleNamespace(
            output_text=json.dumps(_empty_response()),
            model="gpt-5.6-luna-runtime",
            usage=_Usage(),
            output=[],
        )
    )
    analyzer = OpenAIResponsesDocumentAnalyzer(
        settings,
        client=SimpleNamespace(responses=responses),
    )

    result = analyzer.analyze(request)

    assert responses.calls == [request.to_openai_payload()]
    assert responses.calls[0]["store"] is False
    assert responses.calls[0]["tools"] == []
    assert "previous_response_id" not in responses.calls[0]
    assert result.payload == _empty_response()
    assert result.response_model == "gpt-5.6-luna-runtime"
    assert result.usage == {"input_tokens": 12, "output_tokens": 4}
    assert result.refused is False


def test_responses_analyzer_maps_timeout_and_refusal_without_payload_text() -> None:
    settings, request = _request()
    timeout = OpenAIResponsesDocumentAnalyzer(
        settings,
        client=SimpleNamespace(
            responses=_Responses(
                error=openai.APITimeoutError(httpx.Request("POST", "https://example.test"))
            )
        ),
    )

    with pytest.raises(DocumentAnalyzerTimeoutError):
        timeout.analyze(request)

    refusal = OpenAIResponsesDocumentAnalyzer(
        settings,
        client=SimpleNamespace(
            responses=_Responses(
                SimpleNamespace(
                    output_text="",
                    model="gpt-5.6-luna-runtime",
                    usage=_Usage(),
                    output=[
                        SimpleNamespace(
                            content=[SimpleNamespace(type="refusal")]
                        )
                    ],
                )
            )
        ),
    ).analyze(request)

    assert refusal.payload is None
    assert refusal.refused is True
    assert refusal.usage == {"input_tokens": 12, "output_tokens": 4}


def test_invalid_json_reaches_validation_without_exposing_model_text() -> None:
    settings, request = _request()
    analyzer = OpenAIResponsesDocumentAnalyzer(
        settings,
        client=SimpleNamespace(
            responses=_Responses(
                SimpleNamespace(
                    output_text="PRIVATE INVALID MODEL OUTPUT",
                    model="gpt-5.6-luna-runtime",
                    usage=_Usage(),
                    output=[],
                )
            )
        ),
    )

    response = analyzer.analyze(request)

    assert response.payload == "invalid_json"
    assert "PRIVATE" not in repr(response)


@pytest.mark.parametrize(
    ("error", "retryable", "status_class"),
    (
        (
            openai.APIConnectionError(request=httpx.Request("POST", "https://example.test")),
            True,
            None,
        ),
        (
            openai.RateLimitError(
                "limited",
                response=httpx.Response(
                    429,
                    request=httpx.Request("POST", "https://example.test"),
                    headers={"x-request-id": "req-rate"},
                ),
                body=None,
            ),
            True,
            "4xx",
        ),
        (
            openai.InternalServerError(
                "server",
                response=httpx.Response(503, request=httpx.Request("POST", "https://example.test")),
                body=None,
            ),
            True,
            "5xx",
        ),
        (
            openai.BadRequestError(
                "bad request",
                response=httpx.Response(400, request=httpx.Request("POST", "https://example.test")),
                body=None,
            ),
            False,
            "4xx",
        ),
    ),
)
def test_responses_analyzer_classifies_transport_failures_safely(
    error, retryable, status_class
) -> None:
    settings, request = _request()
    analyzer = OpenAIResponsesDocumentAnalyzer(
        settings,
        client=SimpleNamespace(responses=_Responses(error=error)),
    )

    with pytest.raises(DocumentAnalyzerClientError) as captured:
        analyzer.analyze(request)

    assert captured.value.retryable is retryable
    assert captured.value.http_status_class == status_class
    if isinstance(error, openai.RateLimitError):
        assert captured.value.provider_request_id == "req-rate"
    assert "example.test" not in repr(captured.value)


def test_pipeline_runs_each_enabled_cv_in_an_independent_ai_context() -> None:
    settings = AISettings(enabled=True, api_key="test-key")
    analyzer = _Analyzer(
        DocumentAnalyzerResponse(
            payload=_empty_response(),
            response_model="gpt-5.6-luna-runtime",
            usage={"input_tokens": 12, "output_tokens": 4},
        )
    )
    first_text = "First Candidate\nExperienced engineer profile skills"
    second_text = "Second Candidate\nExperienced designer profile skills"
    baseline = analyze_cv_text_result(first_text)

    first = analyze_cv_text_result(
        first_text,
        ai_settings=settings,
        document_analyzer=analyzer,
    )
    second = analyze_cv_text_result(
        second_text,
        ai_settings=settings,
        document_analyzer=analyzer,
    )

    assert first.ai_outcome.status is AIAnalysisStatus.SUCCEEDED
    assert second.ai_outcome.status is AIAnalysisStatus.SUCCEEDED
    assert serialize_report_payload(first.report) == serialize_report_payload(
        baseline.report
    )
    assert len(analyzer.requests) == 2
    first_input = analyzer.requests[0].openai_payload["input"][0]["content"][0]["text"]
    second_input = analyzer.requests[1].openai_payload["input"][0]["content"][0]["text"]
    assert "First Candidate" in first_input
    assert "Experienced engineer profile skills" in first_input
    assert "Second Candidate" not in first_input
    assert "Second Candidate" in second_input
    assert "Experienced designer profile skills" in second_input
    assert "First Candidate" not in second_input
    assert analyzer.requests[0] is not analyzer.requests[1]
    assert all(
        "previous_response_id" not in request.openai_payload
        for request in analyzer.requests
    )


@pytest.mark.parametrize(
    ("analyzer", "failure_reason"),
    (
        (
            _Analyzer(
                DocumentAnalyzerResponse(
                    payload=["invalid"],
                    response_model="gpt-5.6-luna-runtime",
                    usage={},
                )
            ),
            AIFailureReason.INVALID_RESPONSE,
        ),
        (
            _Analyzer(
                DocumentAnalyzerResponse(
                    payload=None,
                    response_model="gpt-5.6-luna-runtime",
                    usage={},
                    refused=True,
                )
            ),
            AIFailureReason.REFUSAL,
        ),
        (
            _Analyzer(error=DocumentAnalyzerTimeoutError()),
            AIFailureReason.TIMEOUT,
        ),
    ),
)
def test_pipeline_failures_are_closed_and_leave_report_bytes_unchanged(
    analyzer,
    failure_reason,
) -> None:
    text = "Candidate Example\nExperienced engineer profile skills"
    baseline = analyze_cv_text_result(text)

    result = analyze_cv_text_result(
        text,
        ai_settings=AISettings(enabled=True, api_key="test-key"),
        document_analyzer=analyzer,
    )

    assert result.ai_outcome.status is AIAnalysisStatus.FAILED
    assert result.ai_outcome.failure_reason is failure_reason
    assert serialize_report_payload(result.report) == serialize_report_payload(
        baseline.report
    )


def test_pipeline_keeps_code_contact_authority_separate_from_ai_semantic_facts() -> None:
    text = (
        "Candidate Example\n"
        "Phone: +49 30 123456\n"
        "Email: candidate@example.com\n"
        "Education\n"
        "Example University\n"
        "Computer Science\n"
        "Experience\n"
        "Example Company\n"
        "Engineer\n"
        "2020 - 2022"
    )
    payload = _empty_response()
    payload["facts"]["education"] = [
                {
                    "kind": "education",
            "institution": {"value": "Example University", "line_ids": ["page-0001-line-0005"]},
            "program": {"value": "Computer Science", "line_ids": ["page-0001-line-0006"]},
            "study_dates": {"value": None, "line_ids": []},
            "status": "present",
        }
    ]
    payload["facts"]["employment"] = [
        {
            "kind": "employment",
            "organization": {"value": "Example Company", "line_ids": ["page-0001-line-0008"]},
            "role": {"value": "Engineer", "line_ids": ["page-0001-line-0009"]},
            "employment_dates": {"value": "2020 - 2022", "line_ids": ["page-0001-line-0010"]},
            "location": {"value": None, "line_ids": []},
            "relationship_type": {"value": "unknown", "line_ids": ["page-0001-line-0009"]},
            "status": "present",
        }
    ]
    result = analyze_cv_text_result(
        text,
        ai_settings=AISettings(enabled=True, api_key="test-key"),
        document_analyzer=_Analyzer(
            DocumentAnalyzerResponse(
                payload=payload,
                response_model="gpt-5.6-luna-runtime",
                usage={},
            )
        ),
    )

    assert result.ai_outcome.status is AIAnalysisStatus.SUCCEEDED
    assert result.ai_outcome.analysis is not None
    assert result.deterministic.facts
    assert all(
        fact.provenance.authority.value == "code"
        for fact in result.deterministic.facts
    )
    ai_facts = result.ai_outcome.analysis.payload["facts"]
    assert ai_facts["education"][0]["authority"] == "ai"
    assert ai_facts["education"][0]["source"] == "document_analyzer"
    assert ai_facts["employment"][0]["authority"] == "ai"
    assert ai_facts["employment"][0]["source"] == "document_analyzer"


def test_enabled_app_defers_ai_until_the_follow_up_request(tmp_path) -> None:
    analyzer = _Analyzer(
        DocumentAnalyzerResponse(
            payload=_empty_response(),
            response_model="gpt-5.6-luna-runtime",
            usage={},
        )
    )
    app = create_app(
        db_path=tmp_path / "ai-wiring.db",
        ai_settings=AISettings(enabled=True, api_key="test-key"),
        document_analyzer=analyzer,
    )
    document = Document()
    document.add_heading("Experience", 1)
    document.add_paragraph("Company: Example Labs Ltd")
    document.add_paragraph("Role: Software Engineer")
    document.add_paragraph("01/2020 - 02/2022")
    hidden = document.add_paragraph().add_run("QUARANTINED_TOKEN")
    hidden.font.hidden = True
    buffer = io.BytesIO()
    document.save(buffer)

    client = TestClient(app)
    response = client.post(
        "/analyze",
        headers={"x-analysis-access-token": "owner-token"},
        files={
            "file": (
                "cv.docx",
                buffer.getvalue(),
                "application/vnd.openxmlformats-officedocument."
                "wordprocessingml.document",
            )
        },
    )

    assert response.status_code == 200
    assert response.json()["ai_analysis"]["status"] == "pending"
    assert len(analyzer.requests) == 0

    retry = client.post(
        f"/analyses/{response.json()['analysis_id']}/ai/retry",
        headers={"x-analysis-access-token": "owner-token"},
    )

    assert retry.status_code == 200
    assert len(analyzer.requests) == 1
    assert retry.json()["ai_analysis"]["status"] == "succeeded"
    provider_input = analyzer.requests[0].openai_payload["input"][0]["content"][0]["text"]
    assert "QUARANTINED_TOKEN" not in provider_input
    assert "█" in provider_input
    assert '"records":' in provider_input
    assert "Example Labs Ltd" in provider_input
